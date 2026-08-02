# -*- coding: utf-8 -*-
"""把「评估指标与评分标准」章节插入 REPORT（原版）.docx，原章节重编号+1，保存为新文件。
数据源：evaluation_report.md（Phase A 两层指标结果）。不修改原文件。"""
import re, shutil, sys, docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'REPORT（原版）.docx'
DST = 'REPORT（原版）_with_metrics.docx'

d = docx.Document(SRC)

# ---------- 工具：在指定段落前插入元素 ----------
def insert_before(ref_para, el):
    ref_para._p.addprevious(el)

def new_para_before(ref_para, text, style_name, bold=False):
    """在 ref_para 前插入一段，返回该段落"""
    p = ref_para.insert_paragraph_before(text, style=style_name)
    if bold:
        for r in p.runs:
            r.bold = True
    return p

def add_run_bold(p, text, bold=True):
    """给段落追加一个加粗 run（用于正文开头关键词加粗）"""
    r = p.add_run(text)
    r.bold = bold
    return r

# 定位：原「2. VAE 类别」标题段落（新章节插它前面）
ref = None
for p in d.paragraphs:
    if p.style.name == 'Heading 1' and p.text.strip().startswith('2. VAE'):
        ref = p
        break
assert ref is not None, '找不到原 2. VAE 类别 标题'
print('插入锚点:', ref.text.strip())

# ============================================================
# 先做原章节重编号 +1（2→3, 3→4, 4→5, 5→6 及所有子标题）
# 注意：重编号在插入新章节之前执行，新插入的 2.x 段落不会被处理
# ============================================================
def renumber(para):
    """对标题段落首编号 +1。EX- 前缀、无编号、摘要/参考文献跳过。
    只修改含前导数字的 run，保留其他 run 格式。"""
    text = para.text.strip()
    if not text:
        return
    if text.startswith('EX-'):
        return
    m = re.match(r'^(\d+)', text)
    if not m:
        return  # 摘要/参考文献等无数字编号
    first = int(m.group(1))
    if first < 2:
        return  # 1. 数据集、Title 不改
    new_first = str(first + 1)
    # 找到第一个含前导数字的 run，只替换其中的数字部分（容忍前导空格）
    replaced = False
    for r in para.runs:
        if re.search(r'\d', r.text):
            r.text = re.sub(r'^(\s*)\d+', r'\g<1>' + new_first, r.text, count=1)
            replaced = True
            break
    if not replaced and para.runs:
        # 兜底：前导数字可能跨 run，直接改第一段 run 文本
        para.runs[0].text = re.sub(r'^(\s*)\d+', r'\g<1>' + new_first, para.runs[0].text, count=1)

for para in d.paragraphs:
    if 'Heading' in (para.style.name or ''):
        renumber(para)

# 修正正文自指 [306]：4.4 节 → 5.4 节
for para in d.paragraphs:
    if 'Heading' not in (para.style.name or '') and '4.4 节' in para.text:
        for r in para.runs:
            if '4.4 节' in r.text:
                r.text = r.text.replace('4.4 节', '5.4 节')

print('重编号完成。锚点现在是:', ref.text.strip())

# 原表格样式（统一 Table Grid）
tbl_style = 'Table Grid'

def make_table(header, rows, widths=None, font_size=None, no_wrap_cols=None):
    """创建表格（不插入，返回 table 元素），复用 Table Grid 样式，表头加粗。
    widths: 每列宽度 twips 列表（1/20 pt）；font_size: 单元格字号 pt；
    no_wrap_cols: 数据行禁换行的列下标集合（防数字被拆成两行，如 0.442→0.44/2）。
    列宽+字号+缩 margin+noWrap 用于防止 13 列大表拆词/换行。"""
    t = d.add_table(rows=1, cols=len(header), style=tbl_style)
    if widths:
        t.autofit = False
        total = sum(widths)
        tblPr = t._tbl.tblPr
        # 表格总宽 tblW
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(total))
        tblW.set(qn('w:type'), 'dxa')
        # 固定布局 tblLayout（防止 Word 重新均分）
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')
        # 缩小单元格内边距（默认 108 twips，压到 40，省出实际列宽）
        cellmar = tblPr.find(qn('w:tblCellMar'))
        if cellmar is None:
            cellmar = OxmlElement('w:tblCellMar')
            tblPr.append(cellmar)
        for side in ('left', 'right'):
            el = cellmar.find(qn('w:' + side))
            if el is None:
                el = OxmlElement('w:' + side)
                cellmar.append(el)
            el.set(qn('w:w'), '40')
            el.set(qn('w:type'), 'dxa')
        # 列网格 tblGrid
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
                gc.set(qn('w:w'), str(w))
        # 每个单元格 tcW
        for j, w in enumerate(widths):
            for cell in t.columns[j].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'dxa')
    # ---- 填充内容（表头加粗；可设字号） ----
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ''
            run = cells[j].paragraphs[0].add_run(str(v))
            if font_size:
                run.font.size = Pt(font_size)
            # 数据列禁用换行（noWrap），防止数字被拆成两行
            if no_wrap_cols and j in no_wrap_cols:
                tcPr = cells[j]._tc.get_or_add_tcPr()
                if tcPr.find(qn('w:noWrap')) is None:
                    tcPr.append(OxmlElement('w:noWrap'))
    return t

# ============================================================
# 新章节内容（数据全部来自 evaluation_report.md）
# ============================================================

# 2.1 为什么需要专门的评估体系
new_para_before(ref, '2. 评估指标与评分标准', 'Heading 1')

new_para_before(ref, '2.1 为什么需要专门的评估体系', 'Heading 2')
p = new_para_before(ref,
    '初期的评估只靠"肉眼 + 一个模糊的分数"（如 80/100、85~90/100），这能快速判断模型好不好，但没法回答"到底好在哪、差在哪"。评审提出两点：评估指标必须齐全（通用指标 + 自设计指标都要算），并且要调研更适合眼底图的专用评分标准。于是我们设计了一套两层评估体系，用同一个特征空间、同一批评估图，对 6 个有 checkpoint 的模型做横向对比。',
    'Normal')
new_para_before(ref, '两层指标的价值互补：通用指标的价值是"可比性"——跟文献里别人报的 FID、KID 对齐；医学意义由专用指标承担——回答"严重症状有没有被保住"。', 'List Bullet')

# 2.2 通用层指标
new_para_before(ref, '2.2 通用层指标（对标文献）', 'Heading 2')
new_para_before(ref, '与生成模型文献对齐的指标，全部用现成库实现，测"生成分布与真实分布的距离"：', 'Normal')
for b in [
    'FID：特征分布距离（Inception 特征，越小越好）—— 最常用的生成质量指标',
    'KID / MMD：分布距离的核方法（KID 无偏、小样本友好，主动强调）',
    'IS：清晰度 × 多样性（越高越好）',
    'Precision / Recall、Density / Coverage：保真度与多样性分解（小样本稳）',
    '1-NN 两样本检验：真图/生成图可否被 kNN 分开（越接近 0.5 越好）',
    'MS-SSIM：生成集内部多样性（越低越好，防模式坍缩）',
    'LPIPS：感知相似度（越低越好）',
    '颜色统计距离：RGB 各通道均值/方差/直方图保真度',
]:
    new_para_before(ref, b, 'List Bullet')

# 2.3 专用层自设计指标
new_para_before(ref, '2.3 专用层自设计指标', 'Heading 2')
new_para_before(ref, '针对眼底图诊断意义设计的指标，通用指标看不到这些医学语义差异：', 'Normal')
for b in [
    '病灶保留率：用颜色阈值在真实图上标定出血（暗红）、渗出（亮黄白）区域，统计生成图里病灶面积占比相对真实图的保留比例 —— 直接量化"病灶有没有被溶解"',
    '血管密度：Frangi 血管增强提取血管，统计血管面积占比分布（无条件模型的血管替代指标）',
    'Vessel Dice：条件模型生成图血管 vs 输入骨架 mask 的重合度（仅条件模型可算）',
    '记忆检测：每张生成图在 Inception 特征空间找最近邻真实图，算 SSIM —— 防止模型"背下"训练图，验证"相似但不相同"',
    'C2ST：小 CNN 真/假二分类交叉验证 AUC，测生成图能否骗过判别器（越低越难识破）',
    'BRISQUE：无参考图像质量（自然图像上训练，眼底图仅作参考）',
]:
    new_para_before(ref, b, 'List Bullet')

# 2.4 指标调研结论
new_para_before(ref, '2.4 指标调研结论（文献依据）', 'Heading 2')
new_para_before(ref, '调研了医学图像生成 / 眼底图专用评估指标的文献，结论是现有两层方案无需推翻，做三处落地增强：', 'Normal')
for b in [
    'FID 小样本有偏：Fréchet 距离用样本均值/协方差计算在样本少时有偏，比较必须同一样本量 —— 我们 330 vs 300 的横向对比有效，绝对值不跨数据集比较；KID 无偏、小样本优先',
    '记忆检测必须独立：文献明确警告 FID/MMD 会随复制率上升而"变好"（模型背题反而得分更高），所以记忆指标与保真指标分开报 —— 正是"相似但不相同"的直接证据',
    'TSTR/TRTR 为终极金标准：合成图用于下游分类任务、对比真实测试集上的表现，才是扩增数据评估的标准范式（对应后续分类器验证）',
    '域适配提取器（RETFound-FD）可选：文献有争议（医学预训练特征不必然更好），列为可选探索，不重算全部历史模型',
]:
    new_para_before(ref, b, 'List Bullet')

# 2.5 评估设置
new_para_before(ref, '2.5 评估设置', 'Heading 2')
for b in [
    '真实基准：330 张真实眼底图，统一 resize 128×128',
    '生成集：每模型 300 张单张图（DDIM 50 步采样，seed 42）',
    '特征空间：InceptionV3（FID/KID/MMD/P-R/D-C/1-NN）、torchvision InceptionV3（IS）、lpips-AlexNet（LPIPS）',
    '评估脚本：eval/metrics_common.py（通用层）+ eval/metrics_fundus.py（专用层）',
]:
    new_para_before(ref, b, 'List Bullet')

# 2.6 通用层评估结果总表
new_para_before(ref, '2.6 通用层评估结果总表', 'Heading 2')
new_para_before(ref, '下表是 6 个有 checkpoint 模型的全套通用层指标（真实 330 张 vs 每模型 300 张）：', 'Normal')
gen_header = ['模型', 'FID↓', 'KID↓', 'MMD↓', 'IS↑', 'P↑', 'R↑', 'D↑', 'C↑', '1-NN(0.5)', 'MS-SSIM↓', 'LPIPS↓', '人工分']
gen_rows = [
    ['FiLM+L1+LPIPS', '178.8', '0.107', '0.001', '2.46', '0.030', '0.442', '0.008', '0.030', '0.984', '0.201', '0.663', '85-90'],
    ['FiLM MSE', '200.2', '0.126', '0.006', '2.26', '0.007', '0.742', '0.001', '0.006', '0.975', '0.187', '0.791', '85'],
    ['条件扩散', '187.2', '0.133', '0.003', '2.34', '0.033', '0.567', '0.009', '0.030', '0.983', '0.251', '0.721', '70'],
    ['基础 DDPM', '188.2', '0.139', '0.003', '2.14', '0.053', '0.473', '0.013', '0.036', '0.979', '0.270', '0.710', '75'],
    ['DCGAN', '229.8', '0.204', '0.007', '1.18', '0.000', '0.000', '0.000', '0.000', '1.000', '0.652', '0.726', '20'],
    ['VAE Large', '188.4', '0.189', '0.004', '2.30', '0.013', '0.000', '0.003', '0.009', '1.000', '0.361', '0.455', '10'],
]
t_gen = make_table(gen_header, gen_rows,
    widths=[1300, 620, 620, 620, 620, 560, 560, 560, 560, 880, 880, 800, 620],
    font_size=8,
    no_wrap_cols=set(range(1, 13)))
insert_before(ref, t_gen._tbl)
new_para_before(ref, '方向：FID/KID/MMD/MS-SSIM/LPIPS 越低越好；IS/Precision/Recall/Density/Coverage 越高越好；1-NN 越接近 0.5 越好。表格中 P/R/D/C = Precision/Recall/Density/Coverage。', 'Normal')

# 2.7 颜色统计结果表
new_para_before(ref, '2.7 颜色统计结果表', 'Heading 2')
new_para_before(ref, 'RGB 各通道统计与真实图 [均值=-0.031, -0.423, -0.707] 的偏差：', 'Normal')
col_header = ['模型', '均值距离↓', '方差距离↓', '直方图距离↓', '生成图均值 [R,G,B]']
col_rows = [
    ['FiLM+L1+LPIPS', '0.287', '0.145', '0.0026', '[0.088, -0.152, -0.234]'],
    ['FiLM MSE', '0.331', '0.153', '0.0029', '[0.086, -0.094, -0.160]'],
    ['条件扩散', '0.379', '0.167', '0.0036', '[0.059, -0.035, -0.047]'],
    ['基础 DDPM', '0.372', '0.168', '0.0035', '[0.065, -0.066, -0.045]'],
    ['DCGAN', '0.483', '0.294', '0.0061', '[0.234, 0.041, 0.011]'],
    ['VAE Large', '0.024', '0.040', '0.0018', '[0.003, -0.412, -0.735]'],
]
t_col = make_table(col_header, col_rows,
    widths=[1500, 1100, 1100, 1200, 2200], font_size=9,
    no_wrap_cols=set(range(1, 5)))
insert_before(ref, t_col._tbl)
new_para_before(ref, '观察：VAE 颜色最贴近真实，但这是"颜色对、结构糊"的假象；所有扩散模型颜色偏亮偏蓝（B 通道偏正），存在系统性色偏，留给后续颜色校正处理。', 'Normal')

# 2.8 专用层自设计指标结果（两张表）
new_para_before(ref, '2.8 专用层自设计指标结果', 'Heading 2')
new_para_before(ref, '病灶 / 血管 / 相似性（Wass = Wasserstein 距离，越小越好；保留率 ≈1 理想，>1 偏多、<1 偏少）：', 'Normal')
spec_header = ['模型', '出血Wass↓', '出血保留', '渗出Wass↓', '渗出保留', '血管Wass↓', 'Vessel Dice↑', '记忆NN-SSIM↓', '复制率↓']
spec_rows = [
    ['FiLM+L1+LPIPS', '0.039', '1.137', '0.317', '13.7', '0.076', '0.180', '0.124', '0.000'],
    ['FiLM MSE', '0.095', '0.254', '0.215', '9.6', '0.083', '0.195', '0.096', '0.000'],
    ['条件扩散', '0.102', '0.199', '0.182', '8.3', '0.049', '0.162', '0.086', '0.000'],
    ['基础 DDPM', '0.105', '0.176', '0.165', '7.6', '0.034', '—', '0.073', '0.000'],
    ['DCGAN', '0.127', '0.000', '0.975', '40.0', '0.040', '—', '-0.010', '0.000'],
    ['VAE Large', '0.068', '0.700', '0.018', '0.30', '0.026', '—', '0.329', '0.000'],
]
t_spec = make_table(spec_header, spec_rows,
    widths=[1500, 850, 850, 850, 850, 850, 850, 1050, 800], font_size=8,
    no_wrap_cols=set(range(1, 9)))
insert_before(ref, t_spec._tbl)

new_para_before(ref, 'C2ST 真伪分类 + BRISQUE 无参考质量：', 'Normal')
c2_header = ['模型', 'C2ST AUC↓', 'BRISQUE↓']
c2_rows = [
    ['FiLM+L1+LPIPS', '0.915', '32.3'],
    ['FiLM MSE', '0.989', '37.8'],
    ['条件扩散', '0.982', '25.0'],
    ['基础 DDPM', '0.980', '18.0'],
    ['DCGAN', '1.000', '36.1'],
    ['VAE Large', '0.999', '32.6'],
]
t_c2 = make_table(c2_header, c2_rows,
    widths=[2200, 1700, 1500], font_size=9,
    no_wrap_cols=set(range(1, 3)))
insert_before(ref, t_c2._tbl)
new_para_before(ref, '真实 BRISQUE = 3.69（BRISQUE 在自然图像上训练，眼底图仅参考）；C2ST = 小 CNN 真/假二分类 5 折交叉验证 AUC（越低越难被识破）。', 'Normal')

# 2.9 无数据可评的模型
new_para_before(ref, '2.9 无数据可评的模型', 'Heading 2')
new_para_before(ref, 'WGAN-GP、StyleGAN2-ADA、VAE Large 1200 三个模型因 checkpoint 权重丢失，无法重新批量生成评估图，本次评估没有数据可评，只能从训练记录给出定性结论：', 'Normal')
for b in [
    'WGAN-GP（标准版 + 削弱版）：判别器/Critic 在 330 张数据上过拟合（分数爆炸/质量不足），训练记录可见，无评估图可评',
    'StyleGAN2-ADA：NVIDIA 官方实现与 PyTorch 2.11 + Windows 环境不兼容，训练 1 tick 后手动中断，无评估图可评',
    'VAE Large 1200：checkpoint 权重丢失（与 800 轮版本同架构），无评估图可评；其 800 轮版本（VAE Large）已纳入评估',
]:
    new_para_before(ref, b, 'List Bullet')

# 2.10 结果解读
new_para_before(ref, '2.10 结果解读', 'Heading 2')
for b in [
    '扩散方法集体胜出：4 个扩散模型 Recall 0.44~0.74 vs GAN/VAE 全 0 —— 验证了项目转向扩散的决策正确',
    'L1+LPIPS 显著抗病灶溶解：出血保留率 1.137（病灶全保留）vs MSE 版 FiLM 0.254（丢失 76%）—— 换损失函数后病灶不再被暗背景拉暗，这是通用指标看不到的医学语义差异',
    'C2ST 与通用层排序自洽：film_l1lpips AUC=0.915 全场最低，DCGAN/VAE 接近 1.0 秒识破',
    '"相似但不相同"强验证：所有生成模型复制率 0%（真实-真实自身 16.4%），最近邻 SSIM 0.07~0.33 远低于真实对自身的 0.548 —— 生成图不复制训练图',
    'DCGAN 无病灶特征：出血保留 0.000、渗出保留 40（全图偏白），判别器过拟合产物完全没有诊断特征',
]:
    new_para_before(ref, b, 'List Bullet')
new_para_before(ref, '综合排名（加权各维度）：1. FiLM+L1+LPIPS（分布最贴近 + 保真多样平衡）；2. FiLM MSE（多样最好）；3. 基础 DDPM / 条件扩散（接近）；4. VAE（颜色假象）；5. DCGAN（全面失败）。与人工评分高度吻合。', 'Normal')

# 保存
d.save(DST)
print(f'已保存: {DST}')
print(f'段落总数: {len(d.paragraphs)}, 表格数: {len(d.tables)}')
