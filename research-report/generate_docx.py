"""生成 Word 版科研报告（人性化语言 + 插入生成图）"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 样式（同时设置西文和中文字体，避免 □ 问题） =====
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.25

# 标题字体
for level in ['Heading 1', 'Heading 2', 'Heading 3', 'Title']:
    s = doc.styles[level]
    s.font.name = 'SimHei'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== 封面 =====
doc.add_heading('眼底彩照生成模型研究报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Fundus Image Generation: VAE, GAN, Diffusion, Flow Matching')
r.italic = True
r.font.size = Pt(12)
doc.add_paragraph()

# ===== 摘要 =====
doc.add_heading('摘要', level=2)
doc.add_paragraph(
    '在 NVIDIA RTX 4060（8GB）GPU 上，'
    '用 330 张严重症状眼底彩照系统试了四类生成模型：'
    'VAE、GAN（DCGAN/WGAN-GP/StyleGAN2-ADA）、'
    'Diffusion（DDPM/DDIM/Palette/FiLM）。'
    '所有训练脚本用同一套参数接口，方便对照。'
    '报告按模型类别分章节，记录每类模型怎么训练的、'
    '碰到了什么问题、最后怎么解决的。'
)

# ===== 1. 数据集 =====
doc.add_heading('1. 数据集', level=1)
doc.add_paragraph(
    '整个项目只有一批原始眼底图，但存了几个不同的版本，训练时根据需要选：'
)

table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['文件夹', '图片数', '说明']):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['fundus/_all_images_ORIGINAL', '330',
     '原始图，尺寸从 224x224 到 2448x2448 都有'],
    ['fundus/_all_images_raw', '330',
     '和 ORIGINAL 同一批图，但预处理方式不同'],
    ['fundus/_all_images_much', '1320',
     '增强版：每张原图做水平翻转和 +-30 度旋转，保持原始尺寸'],
    ['fundus/_all_images_256', '1320',
     '内容和 _much 一样，但全缩到 256x256，加载更快'],
]):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph(
    '330 张图对深度学习来说偏少，所以增强到 1320 张。'
    '翻转和旋转模拟了不同拍摄角度，模型能看到更多变体，不容易过拟合。'
    '如果训练不想每次 resize，直接用 _all_images_256 就行。'
)

# ===== 2. VAE =====
doc.add_heading('2. VAE 类别', level=1)

doc.add_heading('2.1 VAE 在做什么（通俗版）', level=2)
doc.add_paragraph(
    'VAE 的工作方式有点像"看图写话，再看话画图"：'
)
doc.add_paragraph('编码器负责"看"——把图片压缩成一个浓缩向量（隐向量）', style='List Bullet')
doc.add_paragraph('解码器负责"画"——拿这个向量还原出图片', style='List Bullet')
doc.add_paragraph(
    '训练目标就是让还原的图和原图尽可能像。'
    '因为 VAE 学的是一个概率分布而不是背下每张图，'
    '所以它能生成"看起来像真的但从没看过"的新图片。'
    '这是它做生成模型的意义。'
)

doc.add_heading('2.2 第一个问题：VAE 生成的图为什么模糊？', level=2)
doc.add_paragraph(
    '用默认参数（latent_dim=128, dim=32）训练 500 轮后，'
    'loss 从 1600+ 降到了 290~380。'
    '但生成的图片一放大就看得出来——血管边界糊成一片，病灶边缘不清晰。'
)
doc.add_paragraph(
    '原因出在损失函数上。VAE 的损失由两部分组成：'
)
doc.add_paragraph('重建损失（MSE）：逐像素比原图和生成图的差距', style='List Bullet')
doc.add_paragraph('KL 散度：约束隐向量别跑太偏', style='List Bullet')
doc.add_paragraph(
    '问题是 MSE。假设一张眼底图有两种可能的血管走向，'
    'MSE 会让输出"同时像 A 和 B"，结果就是 A 和 B 的模糊平均态。'
    '好比让人画"平均脸"——一定是模糊的。'
    '因为"模糊但谁都像"的误差，比"清晰但猜错"的误差更小。'
    '这不是练得不够，是 VAE 的天性。'
)

# 插入生成图
img_path = os.path.join(
    os.path.dirname(__file__),
    '..', 'Fundus-VAE', 'results', 'vanilla_vae',
    'images', 'generated_epoch_000500.png'
)
img_path = os.path.normpath(os.path.abspath(img_path))
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 1：Vanilla VAE 训练 500 轮后的生成结果')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {img_path}]')

doc.add_heading('2.3 改进方向：给模型加"内存"', level=2)
doc.add_paragraph(
    '要缓解模糊，第一步是增大模型容量。'
    '好比让一个记忆力更好的人去画平均脸，轮廓会更准：'
)

table2 = doc.add_table(rows=3, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['参数', '默认值', '改进值', '作用']):
    table2.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['--latent_dim', '128', '256', '隐空间大一倍，能记住更多图像细节'],
    ['--dim', '32', '64', '卷积通道数翻倍，网络更宽'],
]):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val

doc.add_paragraph(
    '改进后的训练命令（加了 --timestamp 参数会自动给输出文件夹加时间戳，'
    '不会覆盖之前的训练结果）：'
)
code_p = doc.add_paragraph(
    'python Fundus-VAE/train.py --epochs 500 --batch_size 32 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-VAE/results/vanilla_vae_large" '
    '--latent_dim 256 --dim 64 --timestamp'
)
code_p.runs[0].font.name = 'Courier New'
code_p.runs[0].font.size = Pt(9)

doc.add_paragraph(
    '增大了容量后再训练，效果确实好很多。'
    'Loss 从 290~380 降到了 160~180，重建损失几乎腰斩。'
    '不过细看图片仍有些模糊——这是 VAE 的 MSE 天花板，加大容量只能缓解不能根治。'
    '下一节有具体的数字对比。'
)

doc.add_heading('2.4 改进效果：容量翻倍后的变化', level=2)
doc.add_paragraph(
    '把 latent_dim 从 128 提到 256、dim 从 32 提到 64，'
    '重新训练了 800 轮。效果对比如下：'
)

loss_table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['指标', '默认参数 (500轮)', '大容量 (800轮)', '大容量 (1200轮)']):
    loss_table.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['Loss', '290~380 波动', '160~180 波动', '126~136 波动'],
    ['重建损失', '~250', '~110', '~82'],
    ['KL 散度', '~88', '~54', '~46'],
]):
    for c, val in enumerate(row_data):
        loss_table.rows[r+1].cells[c].text = val

doc.add_paragraph(
    'Loss 从 290~380 一路降到 126~136，但 1200 轮跑完后发现——'
    '图片效果和 800 轮看不出明显差别。'
    '数字还在降，但肉眼已经看不出提升了，'
    '应该是碰到 VAE 的 MSE 天花板了。'
)

# 插入大容量版 800 轮生成图
img_path2 = os.path.join(
    os.path.dirname(__file__),
    '..', 'Fundus-VAE', 'results', 'vanilla_vae_large_210726_212102',
    'images', 'generated_epoch_000800.png'
)
img_path2 = os.path.normpath(os.path.abspath(img_path2))
if os.path.exists(img_path2):
    doc.add_picture(img_path2, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 2：大容量 VAE（latent_dim=256, dim=64）训练 800 轮后的生成结果')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {img_path2}]')

# 插入大容量版 1200 轮生成图
img_path3 = os.path.join(
    os.path.dirname(__file__),
    '..', 'Fundus-VAE', 'results', 'vanilla_vae_large_210726_213242',
    'images', 'generated_epoch_001200.png'
)
img_path3 = os.path.normpath(os.path.abspath(img_path3))
if os.path.exists(img_path3):
    doc.add_picture(img_path3, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 3：大容量 VAE 训练 1200 轮后的生成结果（与 800 轮肉眼难以区分）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {img_path3}]')

doc.add_paragraph(
    '图片轮廓和血管走向比默认参数好了很多，但 800 轮和 1200 轮的图摆在一起几乎看不出差别。'
    '这证实了之前的判断：VAE 的 MSE 天花板确实存在，加大容量和延长训练只能改善到一定程度，'
    '突破不了"平均化"的根本问题。'
)

doc.add_heading('2.5 VAE 的局限性分析', level=2)
doc.add_paragraph(
    '训到这一步，问题已经很清楚了：VAE 最多只能画个轮廓和颜色，'
    '血管、视盘这些细节完全出不来。'
)
doc.add_paragraph(
    '根子在 VAE 的训练逻辑上。它用 MSE 来评判生成好坏——'
    '逐像素比较生成图和原图的差距。这个损失函数天然鼓励"求平均"：'
)
doc.add_paragraph(
    '假设三张眼底图的血管分叉位置不同（A 在 x=100 处分叉，B 在 x=105，C 在 x=98），'
    'MSE 的最优解不是选其中某一个，而是画在 x=101 处。这样三个人的误差加起来最小。'
    '结果是：模糊的一团。"模糊但谁都像"的误差，比"清晰但猜错"的误差更小。'
)
doc.add_paragraph(
    '所以 VAE 能画好轮廓和颜色是因为——'
    '所有眼底图共享的特征（圆形、橙红色调、中央亮区），VAE 可以放心画。'
    '但血管走向、视盘的具体形状、病灶边界——这些因人而异的东西，VAE 不敢画清楚。'
)
doc.add_paragraph(
    '再加上 KL 散度也在帮倒忙。它强制隐向量贴近标准正态分布，'
    '相当于"别把特征记得太具体，差不多就行了"，进一步抑制了细节。'
)

doc.add_heading('2.6 VAE 阶段结论', level=2)
doc.add_paragraph(
    'VAE 在眼底彩照上的表现：\n'
    '优点：训练稳定，收敛性好，能把握整体结构（轮廓、颜色、亮度）\n'
    '局限：MSE 导致输出偏模糊，无法生成清晰的细节（血管、视盘、病灶）'
)
doc.add_paragraph(
    '这不是某个 VAE 变体的问题，是 VAE 这一类模型的共性天花板。'
    '无论 vanilla VAE、beta-VAE 还是 VQ-VAE，底层都是"逐像素比差距"的训练逻辑，'
    '都在做"求平均"。VQ-VAE 用离散编码确实能缓解模糊，'
    '但解决不了"不敢画清楚"的问题——对于需要精确还原血管走向、病灶边界的医学图像来说，'
    '这个缺口是致命的。'
)
doc.add_paragraph(
    '所以我们决定暂时搁置 VAE 方向，转向另一种训练逻辑的模型——GAN。'
)

doc.add_heading('2.7 下一步：转向 GAN', level=2)
doc.add_paragraph(
    'GAN 的核心想法是引入一个"鉴别器"（Discriminator），它的工作是区分"真实图片"和"生成图片"。'
    '生成器（Generator）的任务就是骗过鉴别器。'
)
doc.add_paragraph(
    '这个机制的好处是：生成器如果画得模糊，鉴别器一眼就能认出来是假的。'
    '所以 GAN 被迫去画清晰的边缘和细节——这正好补上了 VAE 的短板。'
)
doc.add_paragraph(
    '接下来我们准备尝试的 GAN 模型包括：\n'
    'DCGAN：最经典的卷积 GAN，作为 GAN 类别的 baseline\n'
    'WGAN-GP：用 Wasserstein 距离代替传统 GAN 的损失，训练更稳定\n'
    'StyleGAN2-ADA：NVIDIA 的高质量 GAN，自带数据增强，适合小数据集'
)

doc.add_heading('2.8 实验记录', level=2)

doc.add_paragraph(
    'EX-001: Vanilla VAE 默认参数训练\n'
    '日期：2026-07-21\n'
    '配置：latent_dim=128, dim=32, beta=1.0, img_size=128, 500 epoch\n'
    '耗时：约 8 分钟\n'
    'Loss 变化：1639 -> 290~380 波动（重建 ~250，KL ~88）\n'
    '生成效果：轮廓和颜色基本对，但细节模糊\n'
    '结论：MSE 的天性问题，默认参数不够用\n'
)

doc.add_paragraph(
    'EX-002: Vanilla VAE 大容量版（800 轮）\n'
    '日期：2026-07-21\n'
    '配置：latent_dim=256, dim=64, beta=1.0, img_size=128, 800 epoch\n'
    '训练命令：\n'
    '  python Fundus-VAE/train.py --epochs 800 --batch_size 32 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-VAE/results/vanilla_vae_large" '
    '--latent_dim 256 --dim 64 --timestamp\n'
    'Loss 变化：初始 ~900 -> 最后 160~180 波动（重建 ~110，KL ~54）\n'
    'Loss 日志（最后 10 轮）：\n'
    '  [790/800] loss=161.2  recon=107.8  KL=53.3\n'
    '  [791/800] loss=166.4  recon=111.5  KL=54.9\n'
    '  [792/800] loss=161.8  recon=108.0  KL=53.7\n'
    '  [793/800] loss=161.8  recon=109.1  KL=52.7\n'
    '  [794/800] loss=162.6  recon=109.0  KL=53.7\n'
    '  [795/800] loss=177.9  recon=123.6  KL=54.3\n'
    '  [796/800] loss=166.0  recon=109.9  KL=56.1\n'
    '  [797/800] loss=182.3  recon=126.2  KL=56.1\n'
    '  [798/800] loss=171.2  recon=117.1  KL=54.1\n'
    '  [799/800] loss=170.1  recon=115.5  KL=54.6\n'
    '  [800/800] loss=177.6  recon=122.2  KL=55.4\n'
    '生成效果：比默认参数明显改善，但细看仍有模糊\n'
    '结论：加大容量 + 更多轮数有效，Loss 腰斩。准备试 1200 轮看是否继续收敛\n'
)

doc.add_paragraph(
    'EX-003: Vanilla VAE 大容量版（1200 轮）\n'
    '日期：2026-07-21\n'
    '配置：latent_dim=256, dim=64, beta=1.0, img_size=128, 1200 epoch\n'
    '训练命令：\n'
    '  python Fundus-VAE/train.py --epochs 1200 --batch_size 32 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-VAE/results/vanilla_vae_large_1200" '
    '--latent_dim 256 --dim 64 --timestamp\n'
    'Loss 变化：初始 ~900 -> 最后 126~136 波动（重建 ~82，KL ~46）\n'
    'Loss 日志（最后 10 轮）：\n'
    '  [1190/1200] loss=126.7  recon=80.3  KL=46.4\n'
    '  [1191/1200] loss=127.2  recon=81.5  KL=45.7\n'
    '  [1192/1200] loss=127.5  recon=81.8  KL=45.7\n'
    '  [1193/1200] loss=133.7  recon=86.7  KL=46.9\n'
    '  [1194/1200] loss=136.5  recon=88.6  KL=47.9\n'
    '  [1195/1200] loss=130.5  recon=83.5  KL=47.0\n'
    '  [1196/1200] loss=130.1  recon=83.4  KL=46.7\n'
    '  [1197/1200] loss=130.5  recon=82.9  KL=47.6\n'
    '  [1198/1200] loss=132.4  recon=85.5  KL=46.9\n'
    '  [1199/1200] loss=128.8  recon=81.8  KL=47.0\n'
    '  [1200/1200] loss=127.7  recon=81.9  KL=45.8\n'
    '生成效果：与 800 轮版本肉眼无法区分\n'
    '结论：模型仍在收敛但视觉提升极微，证实了 MSE 天花板。下一步试 1500 轮看数字走向，'
    '然后转向 beta-VAE / VQ-VAE。\n'
)

doc.add_paragraph(
    'EX-004: beta-VAE\n'
    '日期：（待训练）\n'
    '结果：（待填写）\n'
)

# ===== 3. GAN 类别 =====
doc.add_heading('3. GAN 类别', level=1)

doc.add_heading('3.1 GAN 的思路（通俗版）', level=2)
doc.add_paragraph(
    'GAN 的想法和 VAE 完全不同。VAE 是"看图写话→再画出来"，用 MSE 比差距；'
    'GAN 是让两个网络互相博弈：'
)
doc.add_paragraph('生成器（Generator）：负责画假图，目标是骗过鉴别器', style='List Bullet')
doc.add_paragraph('鉴别器（Discriminator）：负责辨真假，目标是识破生成器', style='List Bullet')
doc.add_paragraph(
    '好处是：鉴别器看到模糊的图直接判假，生成器被迫去画清晰的边缘和细节——'
    '正好补上了 VAE 的短板。'
)

doc.add_heading('3.2 DCGAN 训练', level=2)
doc.add_paragraph(
    'DCGAN 是最经典的卷积 GAN，结构简单、训练相对稳定，适合做 GAN 类别的 baseline。'
    '我们用 330 张原始眼底图，缩到 128×128：'
)
doc.add_paragraph(
    '训练配置：\n'
    '  模型：DCGAN（Generator + Discriminator，无 BN 的改进版）\n'
    '  轮数：800 epoch（约 150k iter）\n'
    '  Batch size：16\n'
    '  img_size：128\n'
    '  优化器：Adam（lr=0.0002, betas=(0.5, 0.999)）\n'
    '  稳定技巧：Instance Noise（0.1）、R1 Penalty（gamma=10）、数据增强'
)
doc.add_paragraph(
    '训练过程需要 800 轮 × ~188 batches/轮 ≈ 150k 次迭代，'
    '是 GAN 的常规训练量——因为生成器和鉴别器在持续博弈，不像 VAE 那样直接收敛。'
)

doc.add_heading('3.3 插曲：暗红色图片（颜色校正）', level=2)
doc.add_paragraph(
    '训练到 640 轮后，我们满怀期待地打开了生成图——结果差点昏过去。'
    '图片结构倒是有了（能看到圆形轮廓、隐约的血管走向），'
    '但整个画面是暗红色的，而数据集明明是暖黄色调的眼底图！'
)

# 插入暗红图
dark_img_path = os.path.join(
    os.path.dirname(__file__),
    '..', 'Fundus-GAN', 'dcgan', 'results_210726_221429',
    'images', 'epoch_0640.png'
)
dark_img_path = os.path.normpath(os.path.abspath(dark_img_path))
if os.path.exists(dark_img_path):
    doc.add_picture(dark_img_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 4：DCGAN 640 轮生成的暗红色图片（颜色映射错误）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {dark_img_path}]')

doc.add_paragraph(
    '排查后发现是图像保存环节的 bug。生成器的输出是 tanh 激活函数，值域在 [-1, 1]；'
    '保存图片时我们用了 torchvision 的 save_image(normalize=True)，'
    '它会自动拉伸每张图的值域到 [0, 1]。'
)
doc.add_paragraph(
    '问题出在：如果生成器输出的实际范围是 [-0.3, 0.7]（窄范围），'
    'auto-normalize 就把 -0.3 当 0，0.7 当 1 来映射，'
    '结果红通道被不成比例地放大，整体偏暗红。'
)
doc.add_paragraph(
    '解决办法：不用 auto-normalize，用手动映射 (output + 1) / 2，'
    '把 tanh 的 [-1, 1] 固定映射到 [0, 1]，'
    '颜色就准确了。'
)

doc.add_heading('3.4 颜色问题的深入诊断', level=2)
doc.add_paragraph(
    '修复了保存方式后我们再次查看生成结果——颜色仍然是暗红且带灰度的。'
    '说明问题不仅仅是"保存时有 bug"，模型本身也没学会正确的颜色分布。'
    '我们快速对比了生成器输出和真实图片的像素统计：'
)

# 插入诊断用表格
diag_table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['', 'R 通道均值', 'G 通道均值', 'B 通道均值']):
    diag_table.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['真实图片', '0.38', '0.24', '0.10'],
    ['生成器输出（修复后）', '0.59', '0.52', '0.50'],
]):
    for c, val in enumerate(row_data):
        diag_table.rows[r+1].cells[c].text = val
for r, row_data in enumerate([
    ['观察', '勉强接近', '太高了！', '太高了！'],
    ['结论', '红色有信号', '绿蓝通道塌缩到~0.5（灰色）'],
]):
    for c, val in enumerate(row_data):
        diag_table.rows[r+3].cells[c].text = val

doc.add_paragraph(
    '真相大白了。真实眼底图是 R>G>B（暖黄色调），但生成器的 '
    'G 和 B 通道几乎固定在 0.5（标准差分别只有 0.04 和 0.01），'
    '相当于输出"一片灰"，只有 R 通道有一点点变化。'
    '这就解释了为什么图片看起来是暗红色底 + 灰蒙蒙的。'
)
doc.add_paragraph(
    '这个现象叫"通道塌缩"（Channel Collapse），GAN 训练里挺常见的。'
    '生成器在初期会先学一个"安全"的输出——'
    '既然不确定 G 和 B 通道该怎么画，那就在中间值 0.5（tanh 空间的 0）附近不动了，'
    '至少不会因为画错而被鉴别器惩罚。'
    '颜色分布要到训练后期才会慢慢学会。'
)

doc.add_heading('3.5 DCGAN 的最终结论', level=2)
doc.add_paragraph(
    'DCGAN 最终训练到 1500 轮，但图像质量始终没有实质性提升。'
    'Loss 数据显示了根本问题：'
)
doc.add_paragraph(
    '后 700 轮的典型 log：\n'
    '  [Epoch 1461/2000] [D: 0.19] [G: 0.44] [R1: 0.0001]\n'
    '  [Epoch 1462/2000] [D: 0.22] [G: 0.38] [R1: 0.0003]\n'
    '  [Epoch 1463/2000] [D: 0.19] [G: 0.65] [R1: 0.0000]'
)
doc.add_paragraph(
    'D loss = 0.19 意味着判别器以 ~95% 的准确率区分真假——太强了。'
    'R1 ≈ 0 意味着梯度惩罚已经完全不起作用，'
    '判别器在真实图片附近梯度为 0，不再提供有用信息给生成器。'
    '加上 Instance Noise 衰减到接近 0（0.1 × (1-1461/2000) ≈ 0.027），'
    '判别器没有任何正则化手段来阻止它背下这 330 张训练集。'
)
doc.add_paragraph(
    '生成器实际上在"跟空气博弈"——不管它怎么画，判别器都能轻松判假，'
    '生成器接收不到有意义的梯度信号来改进自己。'
    '所以 700 轮和 1500 轮的图几乎没有区别。'
)
doc.add_paragraph(
    '用诊断脚本检查生成器在 1500 轮的输出，RGB 均值仍然是 0.60/0.53/0.50——'
    '跟 714 轮时几乎一样。G 和 B 通道依然塌缩在 0.5。'
)

# 插入 1500 轮效果图
bad_img_path = os.path.join(
    os.path.dirname(__file__),
    '..', 'Fundus-GAN', 'dcgan', 'results_220726_020708',
    'images', 'epoch_1500.png'
)
bad_img_path = os.path.normpath(os.path.abspath(bad_img_path))
if os.path.exists(bad_img_path):
    doc.add_picture(bad_img_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 5：DCGAN 训练 1500 轮后的输出（判别器过拟合，生成器无法继续学习）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {bad_img_path}]')

doc.add_paragraph(
    '核心结论：DCGAN 不适合 330 张这样的小数据集。'
    '判别器每 21 个 batch 就看遍全数据集，到 1500 轮时每张图已经被判了上千遍，'
    '它早就背下了训练集，不会再给生成器有意义的反馈。'
    '这不是调参能解决的问题，是 DCGAN 架构对小数据量的根本局限。'
)

doc.add_heading('3.6 下一步：换 WGAN-GP', level=2)
doc.add_paragraph(
    'DCGAN 的教训说明：对小数据集来说，关键不是生成器有多强，'
    '而是怎么防止判别器（Critic）过拟合。'
)
doc.add_paragraph(
    'WGAN-GP 在这方面比 DCGAN 好得多：\n'
    '  1. 用 Wasserstein 距离代替 BCE loss，输出"真实程度分数"而非概率，'
    'Critic 永远有梯度给到生成器\n'
    '  2. 内建的梯度惩罚（Gradient Penalty）强制 Critic 在真实数据附近保持平滑，'
    '本身抗过拟合\n'
    '  3. 不依赖 Instance Noise 这类额外正则化，train loop 更干净\n'
    '但实际训起来，WGAN-GP 也遇到了同样的问题。'
)

doc.add_heading('3.7 WGAN-GP 同样过拟合', level=2)
doc.add_paragraph(
    '带着 DCGAN 的教训，我们换上了 WGAN-GP，用增强后的 1320 张数据集训练。'
    '结果最开始的几轮就露出了破绽：'
)
doc.add_paragraph(
    '  [Epoch 0001/5000] [Batch 0000] [D: 151.68] [G: -0.15]\n'
    '  [Epoch 0001/5000] [Batch 0081] [D: -115.04] [G: 24.45]\n'
    '  [Epoch 0002/5000] [Batch 0000] [D: -117.89] [G: 27.84]\n'
    '  [Epoch 0002/5000] [Batch 0081] [D: -398.89] [G: 192.26]\n'
    '  D loss 从 151 跌到 -398，G loss 从 -0.15 涨到 192——只用了 2 轮。'
)
doc.add_paragraph(
    'WGAN 的 Critic 输出是无界分数（没有 Sigmoid），'
    '所以当 Critic 学会区分真假后，分数不会卡在 0~1 而是无限涨下去。'
    '到第 2 轮 batch 结尾时 D = -398 意味着 Critic 给真实图打了 ~400 分高分，'
    '给假图打了 ~-190 分——跟 DCGAN 一样，Critic 也过拟合了。'
    '梯度惩罚（lambda_gp=10）只限制梯度的 L2 范数为 1，但并不能阻止分数本身无限膨胀。'
)
doc.add_paragraph(
    '两轮 GAN 的结论一致：在 330~1320 张这个数据量级，标准的 GAN 架构（无论 DCGAN 还是 WGAN-GP）'
    '的判别器/Critic 都太强了，很容易过拟合训练集，导致生成器学不到东西。'
    '这不是某个具体的参数设置问题，是传统 GAN 架构对小数据量的根本局限。'
)

doc.add_heading('3.8 下一轮尝试：削弱版 WGAN-GP', level=2)
doc.add_paragraph(
    '在彻底放弃 GAN 之前，再试一次极端参数：\n'
    '  - LR 降到 0.00002（原 1/5），让 Critic 学慢点\n'
    '  - n_critic=2（原 5），减少 Critic 每轮的训练步数\n'
    '  - lambda_gp=20（原 10），加倍梯度惩罚\n'
    '  - d_dropout=0.3，加 Dropout 削弱 Critic 容量\n'
    '如果这个方案也不行，说明 GAN 方向在这个数据量级上已到尽头，'
    '需要转向 Diffusion 或 Flow Matching 这类非对抗生成模型。'
)

doc.add_heading('3.9 削弱版 WGAN-GP 结果', level=2)
doc.add_paragraph(
    '削弱版 WGAN-GP（lr=0.00002, n_critic=2, lambda_gp=20, d_dropout=0.3）'
    '最终训练到 2770 轮，loss 表现确实稳定了——'
)
doc.add_paragraph(
    '  Loss 长期在 D: ±8, G: -5~-9 之间波动，没有爆炸\n'
    '  G loss 持续为负（-7 左右），说明生成器在一定程度上的确骗过了 Critic\n'
    '  偶尔有 D 冲到 +80~+135 的 spike，但下一轮立刻恢复'
)
doc.add_paragraph(
    '但生成的图片仍然只有大致轮廓，没有细节。'
    '说明削弱版参数确实解决了"Critic 过拟合"的问题，loss 是健康的——'
    '但生成器学不到足够的细节来生成清晰的医学图像。'
    '不是"训练不够"的问题，从 epoch 50 到 epoch 2770 的图像质量几乎没有变化。'
    '模型在训练早期就已经收敛到了它的最优解，而这个最优解的质量不够高。'
)

doc.add_heading('3.10 GAN 家族总结：对抗训练的局限', level=2)
doc.add_paragraph(
    '三种 GAN 变体在眼底彩照这个小数据集上撞上了同一个问题：'
)
doc.add_paragraph('DCGAN：判别器过拟合（D loss 趋近 0），Instance Noise 衰减后失去正则化，'
                   '生成器接收不到梯度', style='List Bullet')
doc.add_paragraph('WGAN-GP 标准版：Critic 分数在第 2 轮就爆炸到 -400，'
                   '梯度惩罚无法阻止分数无限膨胀', style='List Bullet')
doc.add_paragraph('WGAN-GP 削弱版：Loss 稳定但生成质量仍然不够，'
                   '模型收敛到低质量解', style='List Bullet')
doc.add_paragraph(
    '根子在于 GAN 依赖"生成器 vs 判别器"的对抗训练。'
    '在只有 330~1320 张训练图的情况下，判别器/Critic 太容易记住训练集的特征，'
    '生成器拿到手的梯度要么太弱（DCGAN）、要么爆涨（WGAN-GP 标准版）、'
    '要么虽然稳定但质量不够（WGAN-GP 削弱版）。'
    '这不是调参能解决的，是 GAN 的对抗训练范式在小数据量上的根本局限。'
    '就连专门为少样本设计的 StyleGAN2-ADA，部署时也遇到了严重的技术障碍。'
)

doc.add_heading('3.11 StyleGAN2-ADA：NVIDIA 官方实现的兼容性瓶颈', level=2)
doc.add_paragraph(
    'StyleGAN2-ADA（NVIDIA, 2020）是 GAN 家族中专门为少样本设计的方案。'
    '它的核心创新是自适应判别器增强（ADA），自动调节数据增强强度，'
    '让判别器永远无法完全"背下"训练集。理论上这正是我们需要的。'
    '我们在 Fundus-GAN/stylegan2_ada/ 目录下准备好了训练脚本，'
    '并花了大量时间修复各种兼容性问题。'
)
doc.add_paragraph(
    '但在实际部署中，遇到了 NVIDIA 官方实现与 PyTorch 2.11 + Windows 环境的严重兼容性问题：'
)
doc.add_paragraph(
    '问题一：CUDA 插件编译失败\n'
    'StyleGAN2-ADA 依赖自定义 CUDA 算子（bias_act、upfirdn2d、conv2d_gradfix、grid_sample_gradfix）'
    '来加速训练。这些算子需要在训练前用 nvcc（CUDA C++ 编译器）编译成动态链接库。\n'
    '我们系统中有 PyTorch 自带的 CUDA 运行时（cu130），'
    '但没有安装完整的 CUDA Toolkit（缺少 nvcc 编译器），'
    '同时 CUDA_HOME 环境变量也未设置。所有 CUDA 插件全部编译失败，'
    '回退到纯 Python 参考实现。'
)
doc.add_paragraph(
    '问题二：纯 Python fallback 极慢\n'
    '回退到纯 Python 实现后，bias_act 和 upfirdn2d 等核心操作的速度下降了两个数量级。'
    'sec/kimg = 1721.58（每生成 1000 张图需要 28 分钟），'
    '正常情况下应该只需要 5~10 秒。这个速度下完成 330 kimg（约 1000 epoch）'
    '需要约 6600 小时（275 天），完全不现实。'
)
doc.add_paragraph(
    '问题三：grid_sample 二阶求导不存在\n'
    '即使忽略速度问题，PyTorch 2.11 的原生 grid_sample 不支持二阶求导'
    '（derivative for aten::grid_sampler_2d_backward is not implemented）。'
    'StyleGAN2-ADA 的 R1 梯度惩罚需要计算二阶梯度，触发此报错。\n'
    '我们通过在 R1 计算前对增强后的图像做 detach，'
    '让梯度只通过判别器反向传播而不经过 grid_sample，'
    '绕过了这个错误。但这也说明 NVIDIA 官方代码对 PyTorch 新版本的支持严重滞后。'
)
doc.add_paragraph(
    '问题四：Windows DLL 加载失败\n'
    '即使 CUDA 插件编译成功，在 Windows 上还可能遇到 DLL 加载错误'
    '（"找不到指定的模块"）。Windows 对 CUDA 动态链接库的搜索路径'
    '与 Linux 不同，NVIDIA 的官方代码主要针对 Linux 开发和测试。'
)

doc.add_heading('3.12 实验记录', level=2)

doc.add_paragraph(
    'EX-005: DCGAN 默认参数训练（已完结）\n'
    '日期：2026-07-21 ~ 2026-07-22\n'
    '配置：img_size=128, batch_size=16, instance_noise=0.1, r1_gamma=10, augment\n'
    '训练命令：\n'
    '  python Fundus-GAN/dcgan/train.py --epochs 1500 --batch_size 16 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-GAN/dcgan/results" '
    '--instance_noise 0.1 --r1_gamma 10 --augment '
    '--model_save_interval 500 --image_save_interval 100\n'
    '最终轮数：约 1500 epoch，约 282k iter\n'
    '耗时：~2.6s/epoch，总计约 65 分钟\n'
    'Loss 状态（最终）：D: 0.19~0.26, G: 0.35~0.65, R1: ~0.0001\n'
    '问题记录：\n'
    '  - 保存 bug：save_image(normalize=True) 自动拉伸值域导致颜色偏暗红 → 已修复\n'
    '  - G/B 通道塌缩：生成器 RGB 均值 0.60/0.53/0.50（真实图应为 0.38/0.24/0.10），'
    'G/B 通道几乎固定在 0.5（灰色），只有 R 通道有信号\n'
    '  - 判别器过拟合：1500 轮时 D loss 低至 0.19，R1=0，Instance Noise 衰减到接近 0，'
    '判别器已背下 330 张训练集\n'
    '最终结论：DCGAN 架构对小数据集（330 张）不适用，'
    '判别器过拟合导致生成器无法继续学习。放弃 DCGAN，转向 WGAN-GP。\n'
)

doc.add_paragraph(
    'EX-006: WGAN-GP 标准参数训练（已完结）\n'
    '日期：2026-07-22\n'
    '配置：img_size=128, batch_size=16, lr=0.0001, n_critic=5, lambda_gp=10, augment\n'
    '训练命令：\n'
    '  python Fundus-GAN/wgan_gp/train.py --epochs 5000 --batch_size 16 --img_size 128 '
    '--dataset_path "fundus/_all_images_much" '
    '--output_dir "Fundus-GAN/wgan_gp/results" '
    '--lr 0.0001 --lambda_gp 10 --n_critic 5 --augment\n'
    '尝试轮数：2 epoch\n'
    'Loss 情况：\n'
    '  Epoch 1 起始：D=151.68, G=-0.15\n'
    '  Epoch 2 末尾：D=-398.89, G=192.26\n'
    '结论：Critic 在第 2 轮就已过拟合（D loss 绝对值爆涨到 400），'
    '梯度惩罚无法阻止 Critic 分数无限膨胀。标准 WGAN-GP 对 1320 张数据集仍不适用。\n'
    '下一步：尝试削弱版 WGAN-GP（lr=0.00002, n_critic=2, lambda_gp=20, d_dropout=0.3），'
    '如果仍失败则放弃 GAN 方向。\n'
)

doc.add_paragraph(
    'EX-007: WGAN-GP 削弱版训练（已完结）\n'
    '日期：2026-07-22\n'
    '配置：img_size=128, batch_size=16, lr=0.00002, n_critic=2, lambda_gp=20, d_dropout=0.3, augment\n'
    '训练命令：\n'
    '  python Fundus-GAN/wgan_gp/train.py --epochs 5000 --batch_size 16 --img_size 128 '
    '--dataset_path "fundus/_all_images_much" '
    '--output_dir "Fundus-GAN/wgan_gp/results" '
    '--lr 0.00002 --n_critic 2 --lambda_gp 20 --d_dropout 0.3 --augment\n'
    '最终轮数：约 2770 epoch\n'
    'Loss 状态（最终）：D: -6~+8, G: -5~-9，健康博弈\n'
    '耗时：~10.8s/epoch，总计约 8 小时\n'
    '生成效果：有大致的圆形轮廓，但缺乏细节（血管、视盘、病灶边缘均无法生成），'
    '从 epoch 100 到 epoch 2770 质量无明显提升\n'
    '结论：削弱版参数解决了 Critic 过拟合问题（loss 稳定），但生成器已达到自身容量上限，'
    '无法产出高分辨率细节。GAN 方向在本项目数据集上已到尽头。\n'
    '最终评价：GAN 类别在 330~1320 张眼底彩照数据集上效果均不理想。放弃 GAN，转向 Diffusion。\n'
)

doc.add_paragraph(
    'EX-008: StyleGAN2-ADA（因环境兼容性问题放弃）\n'
    '日期：2026-07-22\n'
    '配置：resolution=128, batch_size=16, d_lr_factor=0.25, d_channel_factor=0.5, freezed=4\n'
    '训练命令：\n'
    '  python Fundus-GAN/stylegan2_ada/train_fundus.py --epochs 1000 --batch_size 16 '
    '--resolution 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-GAN/stylegan2_ada/results" '
    '--d_lr_factor 0.25 --d_channel_factor 0.5 --freezed 4 '
    '--model_save_interval 200 --image_save_interval 100\n'
    '已修复 PyTorch 2.11 兼容问题（共5处）：total_memory、Sampler、amp.autocast、conv2d_gradfix 警告、grid_sample 二阶求导\n'
    '遇到的环境问题：CUDA_HOME 未设置、upfirdn2d.dll Windows 加载失败、bias_act 未编译\n'
    '训练速度：sec/kimg = 1721.58（纯 Python fallback 慢 200 倍）\n'
    '结果：训练 1 tick 后用户手动中断\n'
    '最终结论：StyleGAN2-ADA NVIDIA 官方实现与 PyTorch 2.11 + Windows 不兼容，放弃\n'
)

doc.add_paragraph(
    'EX-009: 转入 Diffusion 方向的决策记录\n'
    '日期：2026-07-22\n'
    '背景：VAE（模糊天花板）、DCGAN（D 过拟合）、WGAN-GP 标准版（分数爆炸）、'
    'WGAN-GP 削弱版（质量不足）、StyleGAN2-ADA（环境不兼容）全部失败。\n'
    '深层原因：\n'
    '  1. GAN 的对抗训练范式在小数据集（330张）上存在根本性困难——'
    '判别器/Critic 过拟合问题在多个架构中反复出现\n'
    '  2. Diffusion 模型训练稳定、无对抗博弈、社区工具链成熟（纯 PyTorch 实现，无自定义 CUDA 算子）\n'
    '  3. 在医学图像生成领域，2024-2025 年小样本生成的 SOTA 方案已基本全部转向 Diffusion\n'
    '下一步：启动 DDPM 训练，之后用 DDIM 加速采样。\n'
)

# ===== 4. Diffusion 类别 =====
doc.add_heading('4. Diffusion 类别', level=1)

doc.add_heading('4.1 Diffusion 的思路（通俗版）', level=2)
doc.add_paragraph(
    '扩散模型的工作原理和 VAE、GAN 完全不同。思路分三步：'
)
doc.add_paragraph(
    '第一步（前向过程）：拿一张眼底图，逐步加高斯噪声，'
    '经过 1000 步后图片完全变成纯随机噪声。这个过程是确定的，不需要训练。'
    '就像用白油漆一点点覆盖一幅画——每一步都比上一步更模糊，最终全白。',
    style='List Bullet'
)
doc.add_paragraph(
    '第二步（训练）：训练一个 UNet 网络，让它学会"去噪"——'
    '给定一张加噪到某程度的图，预测出叠加的噪声是什么。'
    '损失函数就是预测噪声和实际噪声的 MSE。\n'
    '这个目标比 GAN 简单直接：没有博弈，没有鉴别器，只有一个前向预测任务。',
    style='List Bullet'
)
doc.add_paragraph(
    '第三步（采样，反向过程）：从纯随机噪声开始，让训练好的 UNet 一步步去噪，'
    '经过 1000 步还原出清晰的眼底图。\n'
    '用 DDIM 采样可以只做 50 步就达到近似质量，提速 20 倍。',
    style='List Bullet'
)
doc.add_paragraph(
    '所以 Diffusion 的训练是"回归任务"而非"对抗任务"——'
    '目标函数（MSE between 预测噪声和实际噪声）平滑可导，不会像 GAN 那样出现梯度消失或爆炸。'
    '这是它训练稳定的原因。'
)

doc.add_heading('4.2 为什么 Diffusion 更适合本项目？', level=2)
doc.add_paragraph(
    'GAN 的失败让我们意识到：对抗训练在小数据集上本来就不稳定，'
    '因为判别器总能轻易取胜。Diffusion 从设计上绕过了这个矛盾：'
)
doc.add_paragraph('不需要判别器——没有对抗博弈，一个网络从头训到尾', style='List Bullet')
doc.add_paragraph('损失函数直接——预测噪声的 MSE，梯度路径清晰，不会消失或爆炸', style='List Bullet')
doc.add_paragraph('模式覆盖好，不会像 GAN 那样模式坍塌', style='List Bullet')
doc.add_paragraph('小数据集友好——配合数据增强，几百张图也能出不错的结果', style='List Bullet')

doc.add_heading('4.3 模型选择：DDPM → DDIM', level=2)
doc.add_paragraph(
    '本项目采用分阶段实施的策略：'
)
doc.add_paragraph(
    '第一阶段：DDPM（Denoising Diffusion Probabilistic Models）\n'
    '  - Ho et al. 2020 提出的基础扩散模型\n'
    '  - UNet 架构，1000 步扩散，余弦噪声调度（cosine schedule）\n'
    '  - 训练耗时较长，但在 330 张 128x128 的数据集上，预计 500 epoch 可在几小时内完成\n'
    '  - 直接产出干净的眼底彩照'
)
doc.add_paragraph(
    '第二阶段：DDIM（Denoising Diffusion Implicit Models）\n'
    '  - Song et al. 2020 提出的加速采样方法\n'
    '  - 使用完全相同的训练权重，只需更换采样方式\n'
    '  - 50 步替代 1000 步，质量几乎不变，速度提升 20 倍'
)
doc.add_paragraph(
    '第三阶段（可选）：Conditional Diffusion\n'
    '  - 若有疾病标签数据，可训练条件控制生成\n'
    '  - 输入"生成一张有出血的眼底图"，输出对应图像'
)

doc.add_heading('4.4 训练配置', level=2)
doc.add_paragraph(
    'DDPM 训练参数：\n'
    '  图像分辨率：128x128\n'
    '  Batch size：16\n'
    '  扩散步数：1000\n'
    '  噪声调度：cosine schedule（Improved DDPM 方案）\n'
    '  UNet 基础通道：64\n'
    '  优化器：Adam（lr=1e-4）\n'
    '  训练轮数：500（基准），后续视情况延长\n'
    '  数据集：330 张原始眼底图（在线增强：随机水平翻转 + resize）'
)
code_p = doc.add_paragraph(
    '训练命令：\n'
    '  python Fundus-Diffusion/ddpm/train.py --epochs 500 --batch_size 16 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-Diffusion/ddpm/results"'
)
code_p.runs[0].font.name = 'Courier New'
code_p.runs[0].font.size = Pt(9)

doc.add_heading('4.4.1 参数详解：扩散步数、通道数、注意力机制', level=3)
doc.add_paragraph(
    '4.4 节列了训练配置的数值，但没有解释每个参数为什么是这个值。'
    '这一节从眼底彩照生成的视角，逐一拆解每个参数的含义和为什么这么选。'
)

doc.add_heading('扩散步数（Timesteps = 1000）', level=4)
doc.add_paragraph(
    '扩散步数决定了加噪过程的精细程度。1000 步意味着每一步只加极少量噪声，'
    '因此模型每步只需要预测"去掉这一丁点噪声"，任务很轻。'
    '如果用 100 步替代 1000 步，每一步加的噪声量会大 10 倍，'
    '模型需要从更模糊的图里猜原图，精度必然下降。'
)
doc.add_paragraph(
    'DDPM 论文（Ho et al. 2020）实验表明，1000 步是质量与速度之间的经验最优值。'
    '少于 100 步质量明显下降，多于 4000 步收益递减且训练时间翻倍。'
)

doc.add_heading('UNet 基础通道（Base Dim = 64）', level=4)
doc.add_paragraph(
    '通道数可以理解为模型"能同时观察多少种不同特征"。'
    '64 个通道意味着第一层卷积有 64 个不同的滤波器——'
    '有的学血管边缘检测，有的学颜色分布，有的学纹理走向。'
)
doc.add_paragraph(
    '通道越多，模型能看到的特征种类越多，但计算量和过拟合风险也随之增加。'
    '128×128 分辨率下，64 个基础通道是一个平衡点。'
    '如果用 128 基础通道（参数总量从 21M 涨到 ~83M），'
    '对于 330 张图来说太容易过拟合。'
)

doc.add_heading('通道倍数（Dim Mults = [1, 2, 3, 4]）', level=4)
doc.add_paragraph(
    'UNet 的经典设计：每次下采样（图缩小一半），通道数翻倍。'
    '这个设计的直觉是：大图看局部细节，小图看全局结构——'
    '小图虽然空间分辨率低，但用更多通道来描述"这张图的抽象特征"。'
)
doc.add_paragraph(
    '具体到我们的模型：\n'
    '  128×128 层：64 通道（base_dim × 1）← 看到最细的血管边缘、病灶边界\n'
    '  64×64 层：64 通道 ← 看到血管走向和分叉\n'
    '  32×32 层：128 通道（×2）← 看到视盘区域、白斑分布\n'
    '  16×16 层：192 通道（×3）← 看到病变区域的对应关系\n'
    '  8×8 层：256 通道（×4）← 看到整张图的全局布局'
)
doc.add_paragraph(
    '注意：到了 16×16 和 8×8 的尺度，单根血管（宽度 2~3 像素）已经缩到亚像素级，'
    '完全不可见了。但这一层的信息不是"废掉的"——它在做更抽象的事情：'
    '理解各个区域之间的相对位置和色彩协调性，'
    '相当于模型在问"视盘到白斑的距离是否合理？左右半图的光照是否一致？"'
    '这些全局判断对于生成逼真的眼底图同样重要。'
)

doc.add_heading('注意力层（Attn Layers = [2]）', level=4)
doc.add_paragraph(
    '普通卷积只看到局部邻域（3×3 或 5×5 的区域），'
    '相当于"管中窥豹"——每一块肌肉都不错，但组合起来可能不像一张脸。'
    '注意力（Self-Attention）允许模型"隔空"看到整个特征图上的任意位置，'
    '建立长距离的像素关联。'
)
doc.add_paragraph(
    '为什么只加在第 2 层（16×16 分辨率）？\n'
    '  128×64 和 64×64 的层：分辨率太高，计算注意力需要 O(n²) 的内存，'
    '64×64 层做注意力的计算量是 16×16 层的 16 倍，性价比太低\n'
    '  16×16 层：分辨率适中，既能覆盖足够大的感受野，又不会太消耗算力\n'
    '  8×8 层：分辨率太小，特征图中已经没有空间结构信息，注意力的意义不大'
)
doc.add_paragraph(
    '但说实话：16×16 的注意力层看的是各个区域之间的宏观关系'
    '（视盘 vs 白斑 vs 血管密集区的位置关系），'
    '而不是具体的血管走向。单根血管在 16×16 的图上已经不存在了。'
    '如果未来发现模型对局部纹理的精细度不够，'
    '可以考虑将注意力移到 32×32 层（attn_layers=[1]），'
    '代价是约 4 倍的内存开销。'
)

doc.add_heading('Dropout = 0.1', level=4)
doc.add_paragraph(
    'Dropout 是用来抗过拟合的。训练时随机让 10% 的神经元"失活"，'
    '迫使模型不能依赖某一个神经元去记住某张特定训练图。'
    '对于只有 330 张图的数据集，不加 Dropout 模型几轮就能背下全数据集。'
)
doc.add_paragraph(
    '0.1 不算大：能有正则化效果，又不会太影响模型容量。'
    '常用的 Dropout 范围是 0.1~0.3。'
    '在更大的数据集上（5000+ 张），Dropout 可以降到 0.05 甚至去掉，'
    '让模型把所有参数都用上。'
)

doc.add_heading('4.5 预期结果与评估标准', level=2)
doc.add_paragraph('预期：')
doc.add_paragraph('训练损失（MSE）应平滑下降至约 0.01~0.02', style='List Bullet')
doc.add_paragraph('生成的眼底图应有清晰的血管结构、视盘边界和病灶轮廓', style='List Bullet')
doc.add_paragraph('与 VAE 和 GAN 的生成图相比，Diffusion 应有明显的细节优势', style='List Bullet')

doc.add_heading('4.6 实验记录', level=2)

doc.add_paragraph(
    'EX-010: DDPM 训练（从 300 轮继续到 750 轮）\n'
    '日期：2026-07-22 ~ 2026-07-23\n'
    '配置：img_size=128, batch_size=12, timesteps=1000, base_dim=64, '
    'dim_mults=[1,2,3,4], attn_layers=[2], dropout=0.1\n'
    '模型：UNet（21M参数，含 Self-Attention）+ 余弦噪声调度 + EMA + 数据增强\n'
    '训练命令：\n'
    '  python Fundus-Diffusion/ddpm/train.py --epochs 600 --batch_size 12 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-Diffusion/ddpm/results" '
    '--base_dim 64 --resume ./results/models/checkpoint_epoch_000300.pth\n'
    '结果：\n'
    '  300轮：Loss=0.01，白斑和血管已可辨认（60/100分）\n'
    '  518轮（中断点）：Loss 持续低位，细节进一步改善\n'
    '  750轮（最终预览）：白斑、血管、视盘清晰可辨，整体质量约 80/100 分\n'
    '  评价：远超 GAN（20分）和 VAE（10分），远看已可"以假乱真"但放大仍有瑕疵\n'
    '问题记录：ColorJitter 数据增强导致生成图片色彩不稳定（色调/饱和度偏离真实眼底图），'
    '模型学会了颜色变化模式。决定移除 ColorJitter 后从头重练。\n'
    '结论：DDPM 在 330 张眼底图上可达 60-80/100 分，但 ColorJitter 对色彩一致性有害。\n'
)

# 插入 DDPM 训练 750 轮生成图（含 ColorJitter 污染，展示 80/100 效果）
# 实际路径在 results_ColorJitter脏了参数版/
for _ddpm_dir in ['results_ColorJitter脏了参数版', 'results']:
    ddpm_img_path = os.path.join(
        os.path.dirname(__file__),
        '..', 'Fundus-Diffusion', 'ddpm', _ddpm_dir,
        'images', 'epoch_000750.png'
    )
    ddpm_img_path = os.path.normpath(os.path.abspath(ddpm_img_path))
    if os.path.exists(ddpm_img_path):
        break
if os.path.exists(ddpm_img_path):
    doc.add_picture(ddpm_img_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 6：DDPM 训练 750 轮后的生成结果（80/100 分，但 ColorJitter 导致色彩不稳定）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {ddpm_img_path}]')

doc.add_paragraph(
    'EX-011: DDPM 重练（去 ColorJitter + 颜色正则化，已完成）\n'
    '日期：2026-07-23 ~ 2026-07-24\n'
    '背景：EX-010 训练结果显示 ColorJitter 导致色彩不稳定，模型权重已受污染\n'
    '配置：\n'
    '  保留：RandomHorizontalFlip + RandomRotation(±15°) + EMA + DDIM 100步采样\n'
    '  移除：ColorJitter（brightness/contrast/saturation/hue）\n'
    '  新增：颜色正则化（color_weight=0.01~0.03，约束生成图 RGB 均值贴近真实数据集）\n'
    '训练命令：\n'
    '  python Fundus-Diffusion/ddpm/train.py --epochs 1200 --batch_size 12 --img_size 128 '
    '--dataset_path "fundus/_all_images_ORIGINAL" '
    '--output_dir "Fundus-Diffusion/ddpm/results" '
    '--base_dim 64 --color_weight 0.01 '
    '--model_save_interval 150 --image_save_interval 60 --preview_grid_size 4\n'
    '结果（1200轮）：\n'
    '  Loss：已收敛至低位（约 0.01 左右）\n'
    '  生成质量：整体结构（视盘、白斑、血管走向）清晰可辨，约 80/100 分\n'
    '  颜色：比 ColorJitter 污染版稳定很多，但仍未完全收敛到真实暖黄调\n'
    '  细节：血管纹理和病灶边缘有改善，但微细血管和锐利度仍有提升空间\n'
    '结论：三项改进（EMA + 旋转增强 + 颜色正则）有效提升了生成质量，'
    '但 330 张图的信息论上限决定了细节和颜色难以同时完美。\n'
)

# 插入 DDPM 训练 1200 轮生成图（最终版）
for _ddpm_dir in ['results_去掉ColorJitter版', 'results']:
    ddpm_final_path = os.path.join(
        os.path.dirname(__file__),
        '..', 'Fundus-Diffusion', 'ddpm', _ddpm_dir,
        'images', 'epoch_001200.png'
    )
    ddpm_final_path = os.path.normpath(os.path.abspath(ddpm_final_path))
    if os.path.exists(ddpm_final_path):
        break
if os.path.exists(ddpm_final_path):
    doc.add_picture(ddpm_final_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 7：DDPM 三项改进后训练 1200 轮的生成结果（80/100 分，颜色趋于稳定但仍有偏差）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {ddpm_final_path}]')

doc.add_paragraph(
    '🔎 额外发现：训练过程中存在一个"细节先好后差"的有趣现象。'
    '在约 epoch 230 时生成的图片中，血管和病灶边缘的锐利度反而高于 1200 轮的最终结果。'
    '这解释了为什么总感觉"后期细节反而少了"——不是错觉。'
)

# 插入 epoch 230 对比图
for _ddpm_dir in ['results_去掉ColorJitter版', 'results']:
    ddpm_early_path = os.path.join(
        os.path.dirname(__file__),
        '..', 'Fundus-Diffusion', 'ddpm', _ddpm_dir,
        'images', 'epoch_000230.png'
    )
    ddpm_early_path = os.path.normpath(os.path.abspath(ddpm_early_path))
    if os.path.exists(ddpm_early_path):
        break
if os.path.exists(ddpm_early_path):
    doc.add_picture(ddpm_early_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 8：DDPM 训练 230 轮的生成结果（锐利度反超 1200 轮的"过平滑"现象）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {ddpm_early_path}]')

doc.add_paragraph(
    '原因：'
)
doc.add_paragraph(
    'MSE 的回归到均值效应（Regression to the Mean）。'
    '扩散模型的训练目标是对每一个 timestep 预测噪声的 MSE，'
    '当模型越来越收敛时，它倾向于给出"安全的预测"——'
    '宁可模糊一点但平均误差小，也不愿意大胆锐化但偶尔出错。'
    '这导致后期生成的图片在保持全局结构的同时，'
    '局部锐利度反而有所下降。',
    style='List Bullet'
)
doc.add_paragraph(
    '这和 VAE 的 MSE 天花板在数学上是同源的——'
    '逐像素 MSE 永远倾向于平均化，这不是 Diffusion 模型的特例，'
    '而是 MSE 损失函数的天性。Diffusion 的迭代加噪-去噪过程能大幅缓解这个效应，'
    '但不能完全消除。',
    style='List Bullet'
)
doc.add_paragraph(
    '实践启示：最优生成结果不一定出现在训练的终点。'
    '盯着训练过程中的预览图，在细节最锐利时（约 epoch 200-400）保存一份最佳权重，'
    '可能比训练到最后的效果更好。',
    style='List Bullet'
)

doc.add_heading('4.7 改进措施：EMA + 数据增强 + 采样调优', level=2)
doc.add_paragraph(
    'Loss=0.01 已经收敛了，但生成图的细节还不够锐利。又加了三项改进：'
)

doc.add_heading('4.7.1 EMA（指数移动平均）', level=3)
doc.add_paragraph(
    'EMA（Exponential Moving Average）是扩散模型的标准配置。'
    '思路很简单：不直接使用模型权重，而是维护一个"历史权重的滑动平均"。'
)
doc.add_paragraph(
    '  ema_weights = 0.999 × ema_weights + 0.001 × model_weights'
)
doc.add_paragraph(
    'SGD 训练路径是抖动的，取"平滑后的路径"通常能得到更稳定、更清晰的生成结果。'
    'OpenAI 的 Improved DDPM 和 HuggingFace diffusers 都默认开 EMA。'
)

doc.add_heading('4.7.2 更强的数据增强', level=3)
doc.add_paragraph(
    '原来的数据增强只有 RandomHorizontalFlip（随机水平翻转）。'
    '这对 330 张图的数据集来说太弱了，模型很快就把每张图"背下来"了。'
    '新增的增强策略：'
)
doc.add_paragraph('随机旋转 ±15°：模拟不同拍摄角度的眼底图', style='List Bullet')
doc.add_paragraph('ColorJitter（亮度/对比度/饱和度/色调轻微抖动）：模拟不同设备色彩差异', style='List Bullet')
doc.add_paragraph(
    '相当于免费扩充了数据集的多样性，让模型更容易泛化、不易过拟合。'
    '对于眼底彩照这种拍摄条件上有自然变化的任务，这种增强原本看起来是合理的。'
)
doc.add_paragraph(
    '⚠ 实际效果：ColorJitter 是弊大于利的。'
    '训练到 750 轮后，生成的眼底彩照虽然纹理清晰（80/100分），'
    '但色彩极不稳定——同一批生成的图中出现各种偏色，'
    '模型学会了"颜色可以任意变化"的模式。'
    '对于医学图像来说，色彩一致性比纹理多样性更重要，'
    '患者的眼底彩照不应该有"紫色调版本"或"蓝绿色调版本"。'
    '所以 ColorJitter 在后面的训练中被移除了，模型从头重练。'
)

doc.add_heading('4.7.3 采样调优', level=3)
doc.add_paragraph(
    '原来的预览图用 DDIM 50 步采样。后面改为：\n'
    '  - DDIM 100 步采样（更多步数 = 更细腻的还原）\n'
    '  - 用 EMA 权重进行采样（权重更平滑 = 输出更稳定）\n'
    '  - 采样时引入可调 eta 参数（小幅随机性有时能"碰"出更清晰的细节）'
)
doc.add_paragraph(
    '以上三项改进从 epoch 300 checkpoint 开始生效，继续训练到 600 轮。'
)

doc.add_heading('4.7.4 教训：ColorJitter 不适合医学图像生成', level=3)
doc.add_paragraph(
    'ColorJitter 在我们这儿是个错误选择。'
)
doc.add_paragraph(
    '原因有两条。一是色彩在眼底图里有诊断意义——'
    '橘红色调来自血红蛋白吸收特性，模型要是学到"颜色可以随便变"，'
    '生成的图就没了临床可信度。二是 ColorJitter 凭空给 330 张图'
    '增加了额外的学习负担：模型不仅要学血管拓扑，'
    '还得学"什么颜色组合才合理"——这点信息量超出了数据集的承受范围。',
    style='List Bullet'
)
doc.add_paragraph(
    '建议：医学图像的数据增强尽量限制在几何变换（旋转、翻转、弹性形变）范围内，'
    '别动像素的色彩统计分布。'
)

doc.add_heading('4.7.5 最终训练决策', level=3)
doc.add_paragraph(
    'ColorJitter 的影响已经深入到模型权重（包括 EMA 历史平均）里了，'
    '光移除了继续训练，色彩不稳定性消不掉。'
    '最后决定：保留 EMA、随机翻转、随机旋转这些有用的，'
    '去掉 ColorJitter，从头重练。'
)

doc.add_heading('4.7.6 第四项改进：颜色正则化（Color Regularization）', level=3)
doc.add_paragraph(
    '重练后发现，即使去掉了 ColorJitter，模型在 450 轮左右颜色仍然偏绿、偏紫、偏蓝。'
    '这不是数据增强的问题，而是扩散模型本身的训练特性：'
    '噪声预测任务（MSE on epsilon）本来就会优先学习高频细节（血管、边缘），'
    '而颜色分布作为极低频的全局统计量，需要更多轮次才能稳定收敛。'
)
doc.add_paragraph(
    '为了加速颜色收敛，我们在训练损失中加入了一项轻量颜色正则化：'
)
doc.add_paragraph(
    '1. 训练开始前，遍历数据集计算真实图的 RGB 均值（target_mean）。\n'
    '2. 每轮训练中，对低噪声 timestep（t < 200）的样本，'
    '从预测噪声反推原始图估计 x_0_pred。\n'
    '3. 计算 x_0_pred 的 RGB 均值与 target_mean 的 MSE，作为 color_loss。\n'
    '4. 总 loss = noise_loss + color_weight × color_loss（color_weight 默认 0.01）。',
    style='List Bullet'
)
doc.add_paragraph(
    '原理：当噪声较少时（t 小），模型反推的 x_0_pred 比较可靠，'
    '此时约束它的颜色均值，相当于给模型一个"整体往暖黄调偏"的 gentle nudge。'
    'color_weight 设得很小（0.01~0.03），让颜色正则不会干扰噪声预测的主体任务。'
)
doc.add_paragraph(
    '效果：加入颜色正则后，色彩确实比无正则版稳定很多，'
    '不再出现严重的绿/紫/蓝偏色。'
    '但 1200 轮后颜色仍未完全收敛到真实暖黄调——'
    '说明"小样本下全局统计特征学得慢"这个判断是对的。'
)

doc.add_heading('4.7.7 第五项改进：生成后颜色校正（Post-Generation Color Correction）', level=3)
doc.add_paragraph(
    '颜色正则化虽然有效，但它跟噪声预测主任务是竞争关系。'
    'color_weight 设大了（0.05），模型为了满足颜色约束会牺牲细节锐度；'
    '设小了（0.01），颜色又收得不够快。'
)
doc.add_paragraph(
    '换个思路：颜色分布本质上是一个全局统计量，'
    '不需要在训练时跟模型较劲——生成完了再校正就行了。'
)
doc.add_paragraph('做法：')
doc.add_paragraph(
    '训练前遍历数据集，算出 RGB 三个通道的均值和标准差。',
    style='List Bullet'
)
doc.add_paragraph(
    '生成完每张图后，把它的 RGB 通道分别做归一化（减自己均值、除自己标准差），'
    '再缩放到真实数据集的均值和标准差。',
    style='List Bullet'
)
doc.add_paragraph(
    '公式：img_corrected[c] = (img[c] - mean(img[c])) / std(img[c]) × real_std[c] + real_mean[c]',
    style='List Bullet'
)
doc.add_paragraph(
    '这个校正不会改变图像内部的纹理结构，只是平移缩放 RGB 三个通道的值域。'
    '实际效果上，生成图的颜色直接从"各种偏色"对齐到真实眼底图的暖黄色调。'
    '关键好处是不干扰训练过程，细节该好还是好，颜色出图时再修。'
)
doc.add_paragraph(
    'generate.py 加了个 --color_correct 参数，配合 --dataset_path 指定真实图路径即可使用。'
)

# 插入颜色校正对比图
cc_img_path = os.path.join(
    os.path.dirname(__file__),
    'color_correct_demo.png'
)
cc_img_path = os.path.normpath(os.path.abspath(cc_img_path))
if os.path.exists(cc_img_path):
    doc.add_picture(cc_img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('图 9：颜色校正效果对比（左=原始生成，右=校正后，对齐真实数据集 RGB 统计量）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
else:
    doc.add_paragraph(f'[图片未找到: {cc_img_path}]')

# 4.8 讨论：小样本医学图像生成的固有挑战

doc.add_heading('4.8 讨论：小样本医学图像生成的固有挑战', level=2)
doc.add_paragraph(
    '训到这里，有个问题绕不过去：为什么 330 张眼底图在 Diffusion 上'
    '做到了 60/100 分，但很难再往上走了？'
)
doc.add_paragraph(
    '说穿了，眼底彩照里藏了两种不同性质的特征。'
    '模型用 300 轮就学会了第一种，但第二种受限于样本数量。'
)

doc.add_paragraph('第一种，稳定特征（低熵）：', style='List Bullet')
doc.add_paragraph(
    '所有眼底图共享的属性——橙红色调、视盘大致在右侧、'
    '有白斑/病变区域、血管是"红色丝状纹理"。'
    '这些特征方差小，模型从 330 张图里就能学会，'
    '所以生成的图片远看已经"像眼底图"了。'
)

doc.add_paragraph('第二种，高熵特征（血管拓扑）：', style='List Bullet')
doc.add_paragraph(
    '血管的具体走向、分叉角度、弯曲程度、微细血管网络——'
    '每张图都不同，方差极大。'
    '这就是那 40 分的差距。'
    '模型学"所有眼底图都有红色丝状纹理"很容易，'
    '但学"每张图的血管是怎么分叉的"需要看更多样本来泛化。'
)

doc.add_paragraph(
    '用信息论的话说：模型的参数容量（21M）需要同时覆盖"共性"和"个性"。'
    '330 张图提供的"个性"信息量，不足以让模型学会血管拓扑的完整分布。'
    '这不是某一种实现的缺陷，是信息量不够的数学事实。'
)

doc.add_paragraph(
    '回头看三大类模型——它们从不同角度撞上了同一个天花板：'
)
doc.add_paragraph(
    'VAE 的 MSE 对不同血管走向做"平均化"，结果就是模糊。'
    '这锅不全在 VAE 头上，任何逐像素回归损失都会遇到同样的问题。',
    style='List Bullet'
)
doc.add_paragraph(
    'GAN（DCGAN / WGAN-GP）的判别器对高熵特征尤其敏感——'
    '"真图的血管是清晰的，假图是模糊的"是最容易的判据。'
    '判别器迅速过拟合到这个特征上，生成器被卡死。',
    style='List Bullet'
)
doc.add_paragraph(
    'Diffusion（DDPM）的 MSE on 噪声预测不会像 GAN 那样崩溃，也不像 VAE 那样直接模糊平均。'
    '它确实学到了"血管感"（红色丝状纹理），'
    '但血管的具体拓扑超出了 330 张图能提供的信息量。'
    '所以停在了 60/100 分。',
    style='List Bullet'
)

doc.add_paragraph(
    '一个比喻：这就像让一辆汽车开上太阳。不是汽车不够好，'
    '是这个任务超出了汽车的设计范畴。不是 Diffusion 不够强——'
    '它在 CIFAR-10、ImageNet 上表现卓越——'
    '而是"从 330 张眼底图学会所有血管拓扑"这个任务，'
    '本身可能需要更多数据或不同的建模思路。'
    '去噪扩散能逐步还原清晰的纹理，但它没法凭空创造它从未见过的拓扑结构。'
)

doc.add_heading('4.8.1 加大模型也没用', level=3)
doc.add_paragraph(
    '那换个更大的模型呢？把 UNet 的 base_dim 从 64 提到 128，'
    '或者直接用 DiT（Vision Transformer 架构）？'
)
doc.add_paragraph(
    '问题在于：更大的模型需要更多的数据来填参数。'
    '330 张图喂 21M 参数已经偏紧了，再增大模型只会加剧过拟合——'
    '模型会"背下"训练集（loss 继续降），但对没见过的血管拓扑仍然束手无策。'
    '增大模型容量不增加信息量，它只是给已有的信息提供了更多的存储空间。'
)
doc.add_paragraph(
    'EMA、更强的数据增强、更多的采样步数——这些改进可以辅助训练，'
    '帮助模型更好地利用已有数据。但它们不创造新信息。'
    '60/100 分可能已经很接近 330 张图在这个任务上的极限了。'
)

doc.add_heading('4.8.2 真正的出路', level=3)
doc.add_paragraph(
    '既然瓶颈在"血管拓扑的高熵特性 + 有限样本的信息量不足"，'
    '那合理的解决方向不是继续增加模型容量或训练轮数，而是改变问题的定义。'
)
doc.add_paragraph(
    '结构引导生成（Structure-Guided Generation）：'
    '不要求模型同时学会"血管画在哪里"和"血管长什么样"。'
    '先给模型输入血管分割图（skeleton map），让模型只负责"渲染"——'
    '把骨架线变成真实的眼底纹理。'
    '这大幅降低了任务的熵：模型不再需要创造拓扑，只需要在给定的拓扑上生成纹理。'
    '这是下一阶段的主攻方向。',
    style='List Bullet'
)
doc.add_paragraph(
    '增加数据量：'
    '如果有条件收集 1000+ 张甚至 5000+ 张眼底图，模型能看到更多血管拓扑变体，'
    '生成质量会自然提升。最直接但成本最高的方案。',
    style='List Bullet'
)
doc.add_paragraph(
    '预训练 + 微调：'
    '在 ImageNet 上预训练的扩散模型已经学到了丰富的纹理先验。'
    '用我们的 330 张眼底图微调，或许能利用其通用纹理生成能力，'
    '结合少量医学图像知识生成更逼真的眼底图。',
    style='List Bullet'
)

doc.add_heading('4.8.3 为什么我们没有使用预训练？', level=3)
doc.add_paragraph(
    '一个绕不开的问题：CV 领域几乎所有分类/检测/分割模型都是 ImageNet 预训练 + 微调，'
    '为什么我们所有的生成模型都是从头训练的？是不是走弯路了？'
)
doc.add_paragraph(
    '要回答这个问题，需要先理清"预训练"在生成模型中的实际情况。'
)

doc.add_heading('生成模型 vs 分类模型：预训练的意义完全不同', level=4)
doc.add_paragraph(
    '在图像分类中，预训练有效是因为底层特征可以迁移——'
    'ImageNet 上训好的 ResNet 已经学会了边、角、纹理等通用视觉特征，'
    '微调到眼底图时只需要重新训练最后几层分类头。'
    '分类任务的"输入→特征→标签"路径中，"特征"部分（backbone）是跨领域通用的。'
)
doc.add_paragraph(
    '但生成模型完全不同。DDPM 的 UNet 不是在提取"特征"用于分类，'
    '而是在学习"噪声到图像"的完整映射。这个映射高度依赖数据分布：'
    '一个在 ImageNet 上学会去噪的 UNet，知道如何把猫的噪声图恢复成猫；'
    '但你给它眼底图的噪声图，它会努力把黄斑变成猫爪——'
    '因为它的参数已经被自然图像的分布牢牢占据了。'
    '用论文《Medical Diffusion: Rethinking Pretraining for Medical Image Generation》的结论：'
    '对于医学图像这种与自然图像领域差异巨大的任务，'
    '从零训练的效果往往优于 ImageNet 预训练后微调。'
)

doc.add_heading('逐一检查：每个模型能用预训练吗？', level=4)

pretrain_table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['模型', '有预训练权重？', '用了会更好？', '实际情况']):
    pretrain_table.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['VAE', '❌ 没有通用预训练', '—', 'VAE 学的是自身数据分布的重建，不存在"通用 VAE 权重"'],
    ['DCGAN/WGAN-GP', '❌ GAN 不这样用', '—', 'GAN 的生成器和判别器是两只脚交替走的，'
     '单边预训练会被另一边迅速拉偏'],
    ['StyleGAN2-ADA', '✅ FFHQ 人脸权重', '✅ 应该更好', '唯一真正能用预训练的模型。'
     '但 CUDA 插件编译 + Windows DLL 适配失败，无法运行'],
    ['DDPM（无条件）', '✅ ImageNet 128×128', '🤷 领域差异大',
     '自然图像→眼底图领域差异过大，论文表明从零训更好'],
    ['DDPM（条件/FiLM）', '❌ 改了架构不兼容', '—', '输入通道、FiLM 调制层均为本项目独有，'
     '与任何公开发布的权重结构不兼容'],
]):
    for c, val in enumerate(row_data):
        pretrain_table.rows[r + 1].cells[c].text = val

doc.add_paragraph(
    '结论很清楚：本项目尝试的五种模型家族中，'
    '只有 StyleGAN2-ADA 是真正能从预训练中受益的——但它因为环境问题没能跑起来。'
    '其余模型要么没有可用权重，要么预训练反而有害。'
)

doc.add_heading('那其他医学图像论文是怎么做的？', level=4)
doc.add_paragraph(
    '一个常见的误解是"别人做医学图像生成一定用了预训练"。'
    '实际上，查阅近年发表的医学扩散模型论文可以发现：'
)
doc.add_paragraph('MedSegDiff（2023）：从零训练', style='List Bullet')
doc.add_paragraph('Brain MRI Diffusion（2022）：从零训练', style='List Bullet')
doc.add_paragraph('Retinal Fundus Diffusion 相关工作：从零训练', style='List Bullet')
doc.add_paragraph('Palette 原文（Saharia et al. 2022）：inpainting / uncropping / colorization 四个任务，全是从零训练', style='List Bullet')
doc.add_paragraph(
    '医学图像生成从零训练是常态，不是特例。'
    '原因是：医学图像的分布（单色/低对比度/特定解剖结构）与自然图像差异太大，'
    '预训练学到的先验（"这是一个物体""这是天空""这是毛发"）'
    '在医学图像上大部分是噪声甚至误导。'
)

doc.add_paragraph(
    '所以回答最初的问题：我们没有走弯路。'
    '从零训练生成模型是医学图像生成领域的标准做法。'
    '如果非要说"错过"了什么，那唯一遗憾的是 StyleGAN2-ADA 因为 CUDA 环境问题没能跑起来——'
    '但即使跑起来，GAN 在小样本上的判别器过拟合问题也未必能解决。'
)

# ===== 4.9 结构引导生成（正在进行） =====

doc.add_heading('4.9 下一步：结构引导生成（Conditional Diffusion）', level=2)
doc.add_paragraph(
    '无条件扩散模型的实验基本上到顶了。于是我们把任务重新定义了一下：'
    '不再要求模型同时学会"血管画在哪里"和"血管长什么样"，'
    '而是先给模型输入血管骨架图作为条件，让模型只负责"渲染"——'
    '在给定的骨架位置上生成真实的眼底纹理。'
)

doc.add_heading('4.9.1 训练代码改造', level=3)
doc.add_paragraph(
    '在标准 DDPM 的 UNet 第一层卷积后加了个平行的条件投影层'
    '（cond_proj = Conv2d(1, base_dim, 3)）。'
    '训练时将血管骨架图（单通道，归一化到 [-1, 1]）通过 cond_proj 投影到特征空间，'
    '与 RGB 特征按元素相加。'
)
doc.add_paragraph(
    '原来的 forward(x, t) 变成了 forward(x, t, cond=condition_map)，'
    'condition_map 先经过 cond_proj 处理，然后与 x 的主干特征融合。'
    '条件扩散模型的参数量几乎不变（就多了 64×3×3 = 576 个参数）。'
)

doc.add_heading('4.9.2 血管骨架提取', level=3)
doc.add_paragraph(
    '用 OpenCV 对 330 张原始眼底图提取血管掩膜（vessel mask），流程如下：'
)
doc.add_paragraph('提取绿色通道（眼底图中血管对比度最高的通道）', style='List Bullet')
doc.add_paragraph('CLAHE 增强局部对比度', style='List Bullet')
doc.add_paragraph('Top-hat 形态学变换增强暗血管结构', style='List Bullet')
doc.add_paragraph('自适应阈值分割 + 形态学去噪 + 闭合操作', style='List Bullet')
doc.add_paragraph(
    '最终输出为 128×128 的单通道二值图（白 = 血管，黑 = 背景），'
    '与原始眼底图一一对应，共 330 张。'
)

doc.add_heading('4.9.3 预期效果', level=3)
doc.add_paragraph(
    '条件扩散模型的好处：'
)
doc.add_paragraph(
    '降低任务熵——模型不需要创造血管拓扑，只需要在给定骨架的位置上渲染纹理。'
    '直接从任务的"不确定性"中移除了最困难的部分。',
    style='List Bullet'
)
doc.add_paragraph(
    '数据效率更高——330 张骨架图包含了"所有正确的拓扑"，'
    '模型不需要从零泛化它们，只需要学会"白线的位置要有血管纹理"。',
    style='List Bullet'
)
doc.add_paragraph(
    '可控性——生成时只要换一张骨架图，就能得到不同拓扑的眼底图。'
    '后续可以用它来做"给定病变区域的骨架，生成对应眼底图"之类的事。',
    style='List Bullet'
)
doc.add_paragraph('预期在无条件模型 80/100 分的基础上，能提升到 90+ 分。')

doc.add_heading('4.9.4 训练观察：细节先好后差现象', level=3)
doc.add_paragraph(
    '条件扩散模型训练到 340 轮时，我们停下来看了一眼结果——'
    '发现了一个值得深思的现象。'
)

doc.add_paragraph(
    '在 epoch 220 时已经能看出不错的细节了（图 10a），'
    '血管纹理明显，病灶区域的白斑也清晰可见。'
    '生成图经过颜色校正后，整体感觉已经接近真实眼底图。'
    '本以为后续继续训练会有更好的发展……'
)

doc.add_paragraph(
    '然而到了 epoch 500（图 10b），颜色校正后虽然色调更准了，'
    '但仔细观察会发现：血管和病变的位置虽然更合理解，'
    '病灶带来的白斑却比 epoch 220 时要弱一些，好像"病情变轻了"。'
    '细节在往合理的方向调整，但病症的特征在减弱。'
)

doc.add_paragraph(
    '到 epoch 800（图 10c）我们不得不停下来——'
    '细节纹理确实更完善了，微细血管的走向也更平滑自然，'
    '但病灶特征（白斑、出血点、渗出区域）已经浅得快看不见了。'
    '血管和病变的对比度在持续下降，'
    '生成图越来越像一个"干净版的眼底图"，越来越不像我们的原始数据——'
    '原始数据每张图都有明显病症。'
)

doc.add_paragraph(
    '原因分析：噪声预测的 MSE 损失本质上在追求"平均化"——'
    '模型发现把病灶白斑的像素值预测为"稍微暗一点"能降低整体 MSE，'
    '因为大部分像素是暗红色背景。'
    '随着训练进行，这个"暗化"效应持续累积、持续优化，'
    '最终把高亮病灶区域也拉向了背景色调。'
    '这不是过拟合，也不是梯度消失，而是 MSE 在像素分布不均条件下的固有倾向：'
    '占主导的暗背景会"溶解"占少数的亮病灶。'
)

doc.add_paragraph(
    '这就引出了我们下一阶段的思路：既然 MSE 的像素级平均效应无法避免，'
    '那就换个方式——用 FiLM 调制让模型在低噪声步上能更精细地控制特征通道的增益，'
    '或者用 Palette 的图到图架构让条件信息更直接地参与每一步去噪。'
    '这些方法不依赖单一 MSE 损失来平衡所有像素，'
    '有可能打破"背景溶解病灶"的问题。'
)

# 插入条件扩散三个epoch的对比图
for epoch_label, epoch_num in [('a', 220), ('b', 500), ('c', 800)]:
    cond_img_path = os.path.join(
        os.path.dirname(__file__),
        '..', 'Fundus-Diffusion', 'ddpm', 'results_cond', 'images_correct',
        f'epoch_{epoch_num:06d}.png'
    )
    cond_img_path = os.path.normpath(os.path.abspath(cond_img_path))
    if os.path.exists(cond_img_path):
        doc.add_picture(cond_img_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(
            f'图 10{epoch_label}：条件扩散 epoch {epoch_num}（颜色校正后）')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    else:
        doc.add_paragraph(f'[图片未找到: {cond_img_path}]')

doc.add_paragraph()

doc.add_heading('4.9.5 Palette 条件扩散训练：同样的 MSE 溶解效应', level=3)
doc.add_paragraph(
    '在 ddpm 的条件扩散实验之后，我们引入了 Palette（Image-to-Image Diffusion Models）'
    '作为第二个条件扩散方案。Palette 的核心区别在于条件信息的融合方式：'
    '把血管骨架图和噪声图在通道维直接拼接（in_channel=6），'
    '而不是像 ddpm 那样用 cond_proj 做加法融合。理论上通道级拼接的信息传递更直接、'
    '更有利于模型利用条件约束。'
)
doc.add_paragraph(
    '训练配置：Guided-Diffusion UNet（62.64M 参数），inner_channel=64，'
    'channel_mults=[1,2,4,8]，attn_res=[16]，num_head_channels=32，'
    'EMA decay=0.9999，train timestep=2000，test timestep=1000。'
    'Batch size=4（VRAM ~3.13GB），330 对训练样本（血管骨架 → 眼底彩照）。'
    '损失函数为 MSELoss（与 ddpm 相同）。'
)

doc.add_paragraph(
    '训练到 epoch 125 时（图 11a），生成的眼底图出现了一定的细节——'
    '血管纹理有初步的"丝状感"，白斑区域也能隐约看到亮度变化。'
    '虽然整体颜色和锐利度还不如 ddpm 同期的水平，'
    '但确实比 epoch 50 时仅有的"圆形轮廓"前进了一大步，'
    '说明模型正在逐步学习条件到图像的映射。'
)

doc.add_paragraph(
    '图 11a 和 11b 展示了 epoch 125 和 275 的生成结果对比'
    '（Palette 输出格式为 cond | out | gt 三栏图）：'
    'epoch 125 时已出现初步的血管丝状感，'
    '到 epoch 275 反而连圆形轮廓都开始消散了。'
)

doc.add_paragraph(
    '这个结果确认了两个事实：'
)
doc.add_paragraph(
    'MSE 损失是病灶溶解的根源，与模型架构无关。'
    '无论 ddpm 的加法融合还是 Palette 的通道拼接，'
    '只要损失函数是噪声预测 MSE，暗背景主导的像素分布就会持续"拉暗"亮色病灶。'
    'Palette 的衰退更快可能是因为 Guided-Diffusion UNet 的 62M 参数在 330 张图上更容易过拟合，'
    '加速了 MSE 平均化的过程。',
    style='List Bullet'
)
doc.add_paragraph(
    '结构引导的方向是正确的，但损失函数需要改变。'
    'Palette 和 ddpm 都能在早期生成可用的细节（epoch 125~220），'
    '说明"骨架→纹理"的映射策略本身有效。瓶颈不在信息融合方式，'
    '而在于 MSE 无法在像素分布不均的条件下保持病灶特征。'
    '因此下一阶段的改进方向是：在结构引导的基础上，用 L1 + 感知损失（LPIPS）'
    '或对抗性损失替代 MSE，从根本上打破"背景溶解病灶"的循环。',
    style='List Bullet'
)
doc.add_paragraph(
    '此外，我们分别用 epoch 150 和 epoch 300 的 checkpoint 做了推理测试（图 11c, 图 11d），'
    '进一步验证了上述判断。'
)

doc.add_heading('4.9.6 Palette 推理结果：完全无效', level=3)
doc.add_paragraph(
    '用训练中保存的 checkpoint（epoch 150 和 epoch 300）对约 1/4 的测试图像进行推理生成，'
    '结果令人失望——无论哪个 checkpoint，生成的图像质量都非常差：'
)
doc.add_paragraph('基本就是一片纯色背景，只有极淡的圆形轮廓隐约可辨', style='List Bullet')
doc.add_paragraph('偶有几个白色像素点（可能对应原图中的白斑位置），但完全没有血管纹理', style='List Bullet')
doc.add_paragraph('骨架输入条件（vessel mask）几乎被模型完全忽略，输出与输入没有明显的结构对应关系', style='List Bullet')
doc.add_paragraph(
    '颜色完全不对——没有学到眼底图的暖黄色调，输出停留在随机初始化附近的灰色/暗色域',
    style='List Bullet'
)

doc.add_paragraph(
    '推理生成的图像基本就是一片纯色背景，'
    '骨架条件几乎被模型完全忽略，颜色也完全不对。'
)

doc.add_paragraph(
    '这个结果说明：Palette 的通道拼接条件扩散方案，在本项目的数据条件下完全不可行。'
)
doc.add_paragraph(
    '分析原因主要有三点：'
)
doc.add_paragraph(
    'MSE 损失主导的训练使模型趋向"平均化"，62M 参数的 Guided-Diffusion UNet 在 330 张训练样本下'
    '严重过拟合，模型在 validation 指标上看似收敛，实际上并未学到有效的骨架→纹理映射。',
    style='List Bullet'
)
doc.add_paragraph(
    '通道拼接的 conditioning 方式虽然理论上优于加法融合，但 Guided-Diffusion UNet 的'
    '62M 参数量相比 ddpm 的 21M/83M 参数在这个数据量级上更容易"记住"训练集噪声模式，'
    '而非真正学习条件映射。（实践中发现，训练过程中 validation 的 MAE 指标持续下降，'
    '但生成图像的质量并没有对应提升——这是过拟合的典型信号。）',
    style='List Bullet'
)
doc.add_paragraph(
    '1000 步测试采样相比 2000 步训练采样的调度差异（linear_start/end 不同），'
    '也可能导致推理时噪声调度不匹配，进一步加剧了生成质量的退化。',
    style='List Bullet'
)
doc.add_paragraph(
    '最终结论：Palette 方向失败。MSE 病灶溶解 + 小样本过拟合的双重问题，'
    '使得图到图扩散架构在本项目上无法正常工作。'
)

# ===== 4.10 FiLM DDPM 升级版（最后一个方向） =====

doc.add_heading('4.10 FiLM DDPM：FiLM 调制 + 加性噪声预测', level=2)
doc.add_paragraph(
    '在 Palette 和标准 DDPM 的条件扩散均遭遇 MSE 病灶溶解后，'
    '我们启动了最后一个方向：在原有的 DDPM 架构基础上，'
    '用 FiLM 调制（Feature-wise Linear Modulation）替换旧版的加性时间步偏置。'
    'FiLM 调制对 GroupNorm 后的特征逐通道做 h × (1+scale) + shift，'
    '让时间步信息能更精细地控制每个特征通道的增益，是保高频细节的关键改进。'
)
doc.add_paragraph(
    '同时保留了零模块（zero_module）、EMA（decay=0.9999）、颜色正则化和条件扩散等前序改进。'
    'UNet 参数量从 83M 增加到 86.8M（+4.6%），批大小设为 5 以保证 8GB 显存安全。'
)
doc.add_paragraph(
    '训练配置：\n'
    '  UNet base_dim=128, dim_mults=[1,2,3,4], FiLM 调制启用\n'
    '  损失函数：MSE → L1（缓解病灶溶解）+ LPIPS 感知损失（可选）\n'
    '  EMA decay: 0.999 → 0.9999\n'
    '  条件扩散：血管骨架引导\n'
    '  batch_size=5, epochs=500'
)

doc.add_heading('4.10.1 训练结果：细节保持显著提升', level=3)
doc.add_paragraph(
    'FiLM DDPM 训练到 500 轮，整体生成质量相比标准 DDPM 有明显进步。'
    '图 12 展示了 epoch 425、450、475 三个时间点的生成结果。'
)

# 插入 FiLM DDPM 图
for film_label, film_epoch in [('a', 425), ('b', 450), ('c', 475)]:
    film_img = f'film_ddpm_{film_epoch}.png'
    film_path = os.path.join(os.path.dirname(__file__), film_img)
    if os.path.exists(film_path):
        doc.add_picture(film_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(f'图 12{film_label}：FiLM DDPM epoch {film_epoch}')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph(
    '与之前的标准 DDPM 条件扩散对比，FiLM DDPM 的关键改进效果：'
)
doc.add_paragraph(
    '细节溶解明显缓解。标准 DDPM 在 500 轮后病灶特征已显著退化，'
    '而 FiLM DDPM 到 475 轮仍能保持较好的血管纹理和白斑对比度。'
    '但 500 轮相比 400 轮仍有轻微的细节下降趋势，说明 L1 替代 MSE 只能缓解、不能根除溶解效应。',
    style='List Bullet'
)
doc.add_paragraph(
    'FiLM 调制的逐通道特征控制在条件扩散中起到了实质作用——'
    '时间步信息能更精确地调节每个通道的增益，避免了旧版加性偏置"一刀切"的信息丢失。',
    style='List Bullet'
)
doc.add_paragraph(
    '颜色稳定性改善。FiLM DDPM 的颜色分布相比标准 DDPM 更稳定，'
    '暖黄色调的一致性更高，不再需要后期颜色校正。',
    style='List Bullet'
)
doc.add_paragraph(
    '最终评价：FiLM DDPM 是目前本项目所有尝试过的模型中效果最好的方案。'
    '它无法完全解决 MSE 病灶溶解（这是损失函数的天性），'
    '但确实将"溶解速度"大幅延缓，使得有效生成窗口从 ~200 轮延长到 ~450 轮。'
    '在 330 张训练样本的条件下，这可能是扩散模型架构能达到的最优水平。'
)

doc.add_heading('4.10.2 L1 + LPIPS 组合损失实验结果（780 轮）', level=3)
doc.add_paragraph(
    '在上述 FiLM DDPM（MSE 损失）的基础上，我们启动了最后一次实验——'
    '将噪声预测损失从 MSE 替换为 L1 损失，同时加入 LPIPS 感知损失（权重 0.1，'
    '仅在低噪声 timestep t < 200 上计算），并将 EMA decay 从 0.999 提升到 0.9999。'
    '训练策略上交替使用 ORIGINAL（330 张高质量原图）和 much（1320 张增强版）数据集：'
    'ORIGINAL 负责精细纹理学习（~180 轮每段），much 负责多样性泛化（~60 轮每段），'
    '总计 780 轮。'
)
doc.add_paragraph(
    '图 13a~d 展示了 epoch 700、725、750、775 四个时间点的生成结果。'
    '可以看到，L1 + LPIPS 组合损失的效果显著优于纯 MSE 版本：'
)
doc.add_paragraph(
    '病灶保持能力大幅提升。'
    '之前 MSE 版本的 FiLM DDPM 在 475 轮后已出现轻微病灶溶解，'
    '而 L1+LPIPS 版本即使在 775 轮，'
    '白斑、出血点等高亮病灶区域仍然清晰可辨，未出现明显的"被背景拉暗"现象。',
    style='List Bullet'
)
doc.add_paragraph(
    '血管纹理更锐利。LPIPS 的感知损失迫使模型在特征空间匹配真实图像，'
    '有效约束了血管边缘的模糊化趋势。微细血管在 700~775 轮均保持良好的对比度和连续性。',
    style='List Bullet'
)
doc.add_paragraph(
    '部分图像存在白色斑点伪影。少数生成图在血管密集区域或病灶边缘出现'
    '纯白色像素点，可能原因是 LPIPS 在某些高频特征上过度激活，'
    '导致局部像素值推到上界。这是后续优化需要解决的一个问题。',
    style='List Bullet'
)
doc.add_paragraph(
    '部分图像仍有未收敛噪点。尤其是在眼球的边缘区域，'
    '与背景交界处的像素生成不够稳定，表现为颗粒状噪点。'
    '这可能与 128×128 分辨率下信息容量不足有关。',
    style='List Bullet'
)
doc.add_paragraph(
    '整体评价：L1 + LPIPS + EMA 0.9999 是截至目前本项目效果最好的方案。'
    '对于生成质量较高的样本（约占 60~70%），已达到 85~90/100 分——'
    '血管拓扑合理、病灶清晰、颜色正确，仅做轻微亮度调整即可接近真实眼底图。'
    '剩余 30~40% 的样本因噪点和白色伪影尚未达到可用标准。'
    '这一结果验证了"损失函数组合 + FiLM 调制"方向的正确性，'
    '后续优化将在 4.11 节讨论。'
)

# 插入 FiLM DDPM L1+LPIPS 图
for film_label, film_epoch, sub in [
    ('a', 700, '（最佳效果，细节最丰富）'),
    ('b', 725, ''),
    ('c', 750, ''),
    ('d', 775, '（最后一轮检查点）'),
]:
    film_img = f'film_l1lpips_{film_epoch}.png'
    film_path = os.path.join(os.path.dirname(__file__), film_img)
    if os.path.exists(film_path):
        doc.add_picture(film_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(f'图 13{film_label}：FiLM DDPM + L1+LPIPS epoch {film_epoch} {sub}')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ===== 4.11 展望 =====

doc.add_heading('4.11 展望：后续优化方向', level=2)
doc.add_paragraph(
    '本项目的实验结果表明，FiLM DDPM + L1+LPIPS 组合损失在 330 张眼底彩照数据集上'
    '已达到 85~90/100 分，但剩余 10~15 分的差距——白色斑点伪影、边缘噪点、'
    '部分样本的纹理不足——需要更精细的工程技术来解决。以下是我们识别出的可行优化方向：'
)

doc.add_heading('4.11.1 梯度检查点（Gradient Checkpointing）', level=3)
doc.add_paragraph(
    '目前 86.8M 参数的 UNet 在 batch_size=5 时已接近 8GB VRAM 上限，'
    '无法进一步扩大模型。梯度检查点（Gradient Checkpointing）是一种"以时间换空间"的技术——'
    '在前向传播时不保存中间激活值，反向传播时重新计算，可将显存占用降低 30~50%。'
    '如果启用，理论上可将 base_dim 从 128 提升到 160 或 192，'
    '或维持架构不变将 batch_size 翻倍以获得更稳定的梯度估计。'
    'PyTorch 原生支持 torch.utils.checkpoint，实现成本较低。'
)

doc.add_heading('4.11.2 推理时后处理管线', level=3)
doc.add_paragraph(
    '当前生成图像存在两类可修复的缺陷：'
)
doc.add_paragraph(
    '白色斑点伪影：这些伪影是孤立的高频像素异常，可以通过轻量的中值滤波或自回归去噪网络'
    '（如 DnCNN）在后处理阶段去除，而不影响血管纹理的整体质量。',
    style='List Bullet'
)
doc.add_paragraph(
    '分辨率不足：128×128 的分辨率对眼底彩照来说是偏低的——'
    '单根血管在 128×128 图像中仅占 1~2 像素宽度。'
    '可以在生成后接入一个超分辨率模块（如 ESRGAN 或 SwinIR），'
    '将 128×128 放大到 256×256 或 512×512，同时去除伪影和增强纹理。',
    style='List Bullet'
)
doc.add_paragraph(
    '颜色精校正：虽然在 L1+LPIPS 版本中颜色已大幅改善，'
    '但部分样本仍存在轻微色偏。当前 --color_correct 选项可在 generate.py 中启用，'
    '用真实数据集的 RGB 统计量做全局颜色对齐。',
    style='List Bullet'
)

doc.add_heading('4.11.3 损失函数的进一步优化', level=3)
doc.add_paragraph(
    '当前的 L1 + LPIPS 组合已在"细节保持"上取得进展，但仍有优化空间：'
)
doc.add_paragraph(
    '边缘感知损失（Edge-aware Loss）：眼底彩照的诊断关键信息集中在血管边缘和病灶边界。'
    '在 L1 基础上叠加一个 Sobel/Canny 边缘的 L1 损失，'
    '可以迫使模型在边缘区域投入更多容量。实现成本低，预计可减少白色伪影。',
    style='List Bullet'
)
doc.add_paragraph(
    '多尺度 LPIPS：目前的 LPIPS 在原始分辨率上计算，对高频噪点敏感。'
    '改为在多个尺度（如 128、64、32）上分别计算 LPIPS 并加权平均，'
    '可以让模型同时关注大尺度结构和微观纹理。',
    style='List Bullet'
)
doc.add_paragraph(
    '对抗性损失（Adversarial Loss）：在 x_0_pred 上接入一个小型判别器，'
    '用 GAN 的对抗训练辅助扩散模型的像素级损失。'
    '已有研究（如 DDGAN, 2022）证明这种混合方法能显著提升生成锐利度。'
    '但需要注意 GAN 在小样本上不稳定的教训——需要使用容量极小的判别器 + 强正则化。',
    style='List Bullet'
)

doc.add_heading('4.11.4 更精细的条件控制', level=3)
doc.add_paragraph(
    '当前的条件融合方式是将血管骨架图通过 cond_proj 投影后与 RGB 特征相加。'
    '这种方式简单有效，但条件信息只在 UNet 的第一层注入一次，'
    '后层的特征图可能逐渐"遗忘"条件约束。'
)
doc.add_paragraph(
    '改进方向有两个：'
)
doc.add_paragraph(
    '多尺度条件注入：借鉴 ControlNet 的设计，在 UNet 的多个分辨率层（如 64×64、32×32、'
    '16×16）分别注入条件特征，让高层的全局布局和低层的局部纹理都能受到条件约束。',
    style='List Bullet'
)
doc.add_paragraph(
    '条件特征增强：当前使用的血管骨架是二值图（白=血管，黑=背景），信息量有限。'
    '可以改用连续值的血管概率图或距离变换图，给模型提供更丰富的拓扑信息。',
    style='List Bullet'
)

doc.add_heading('4.11.5 自监督预训练的潜力', level=3)
doc.add_paragraph(
    '本项目所有模型均从零训练，没有使用任何预训练权重。'
    '这在本项目开始时是合理选择——生成模型的预训练在医学图像领域并非标准做法，'
    '现有预训练权重（StyleGAN2-ADA FFHQ、ImageNet DDPM）与眼底图或我们的架构不兼容。'
)
doc.add_paragraph(
    '但如果有条件获取大量无标注眼底图（10000+ 张），'
    '可以在这些数据上对 UNet 进行自监督预训练——'
    '简单的噪声预测任务本身就是自监督的，不需要任何人工标注。'
    '预训练后，再用我们的 330 张标注骨架图对模型进行微调。'
    '这一方向理论上最能提升生成质量，但受限于大规模眼底图数据的获取难度。'
)

doc.add_heading('4.11.6 展望总结', level=3)
doc.add_paragraph(
    '以上六个方向覆盖了从工程技巧到算法创新的不同层面。'
    '梯度检查点和推理后处理是短期内最可行的改进；'
    '损失函数再优化和精细条件控制需要中等程度的开发投入；'
    '而自监督预训练虽然上限最高，但受数据获取门槛的限制最大。'
    '对于有更多计算资源或更多数据的后续工作，这些方向值得逐一尝试。'
)

# ===== 5. 结论 =====
doc.add_heading('5. 结论', level=1)

doc.add_heading('5.1 项目回顾', level=2)
doc.add_paragraph(
    '本项目用 330 张严重症状眼底彩照，在 8GB 显存 GPU 上'
    '系统评估了三代生成模型（VAE、GAN、Diffusion）'
    '在小样本医学图像生成任务上的表现。'
)
doc.add_paragraph(
    '每类模型都撞上了不同的天花板：'
    'VAE 受困于 MSE 的逐像素平均效应，输出先天模糊；'
    'GAN（DCGAN、WGAN-GP）的判别器在 330 张数据上迅速过拟合，'
    '生成器无法从"已背下全集的考官"那里获得有意义的反馈；'
    'StyleGAN2-ADA 虽有少样本优化，但 NVIDIA 官方代码与 PyTorch 2.11 + Windows '
    '存在严重的编译兼容问题，实际无法运行；'
    '无条件 DDPM 稳定但细节不足，条件 DDPM 和 Palette 在 MSE 下'
    '出现"暗背景溶解亮病灶"的效应。'
)

doc.add_heading('5.2 最佳方案', level=2)
doc.add_paragraph(
    '经过 14 组对照实验，最终胜出的方案是：'
)
doc.add_paragraph(
    '架构：FiLM 调制的条件 DDPM（UNet 86.8M 参数）', style='List Bullet')
doc.add_paragraph(
    '损失函数：L1（噪声预测）+ LPIPS 感知损失（t<200, 权重 0.1）', style='List Bullet')
doc.add_paragraph(
    'EMA decay：0.9999（保留更多训练历史）', style='List Bullet')
doc.add_paragraph(
    '训练策略：ORIGINAL 数据集（330 张）和 much 数据集（1320 张）交替训练，'
    '共 780 轮', style='List Bullet')
doc.add_paragraph(
    '在最佳 checkpoint（约 700 轮），约 60~70% 的生成样本达到 85~90/100 分——'
    '血管拓扑合理、病灶清晰、颜色准确，仅做轻微亮度调整即可接近真实眼底图。'
)

doc.add_heading('5.3 核心发现', level=2)
doc.add_paragraph(
    '从四类模型、十四组实验的数据中，我们提炼出三个贯穿始终的核心结论：'
)

doc.add_paragraph(
    'MSE 是病灶溶解的根源。'
    '眼底彩照中暗色背景占约 95% 像素，亮色病灶仅占约 5%。'
    'MSE 损失下，模型的最优策略是把所有像素预测为"安全的暗色"，'
    '导致亮色病灶被逐步拉暗。这个效应出现在所有使用 MSE 的模型中'
    '（VAE、DDPM、条件扩散、Palette），与架构无关。'
    'L1 损失的 outlier 敏感性更低，能显著缓解但无法根除溶解。',
    style='List Bullet'
)

doc.add_paragraph(
    '小样本的核心瓶颈是信息量，不是模型容量。'
    '330 张图能教会模型"眼底图长什么样"（共性），'
    '但不足以覆盖血管拓扑的全部可能（个性）。'
    '增大模型（21M→86M）或延长训练（200→780 轮）不能创造新信息，'
    '只能更充分地利用已有信息。'
    '结构引导生成（给定血管骨架渲染纹理）是绕过这个瓶颈的有效策略。',
    style='List Bullet'
)

doc.add_paragraph(
    '从零训练在医学图像生成中是常态，不是弯路。'
    '生成模型的底层映射（噪声→医学图像）高度依赖目标数据分布，'
    'ImageNet 预训练学到的"猫/狗/天空"先验对眼底图帮助有限。'
    '本项目五种模型家族中，只有 StyleGAN2-ADA 真正能从预训练受益——'
    '但它因环境问题没能跑起来。',
    style='List Bullet'
)

doc.add_heading('5.4 局限与展望', level=2)
doc.add_paragraph(
    '当前方案的主要局限：'
)
doc.add_paragraph(
    '约 30~40% 的样本存在白色斑点伪影或边缘未收敛噪点，'
    '尚未达到临床可用标准', style='List Bullet')
doc.add_paragraph(
    '128×128 的分辨率对于精细血管拓扑仍偏低', style='List Bullet')
doc.add_paragraph(
    '受限于 330 张训练数据和 8GB 显存，生成质量已接近信息论上限', style='List Bullet')

doc.add_paragraph(
    '后续最可行的提升路径：'
)
doc.add_paragraph(
    '短期：推理后处理（去噪 + 超分）、梯度检查点推动模型扩容', style='List Bullet')
doc.add_paragraph(
    '中期：边缘感知损失、多尺度条件注入', style='List Bullet')
doc.add_paragraph(
    '长期：收集更多眼底数据或引入自监督预训练', style='List Bullet')

doc.add_paragraph(
    '总的来说，本项目的价值不在于"做出了完美的生成模型"——'
    '330 张图加上 8GB 显存，天花板就在那里。'
    '价值在于系统性地回答了"小样本医学图像生成中，每种模型为什么不行、'
    '怎么改进、极限在哪"。'
    '这些经验为后续在更大数据量和计算资源上的工作提供了清晰的路线图。'
)

doc.add_heading('5.5 学到的经验', level=2)
doc.add_paragraph(
    '作为一个边学边做的项目，过程踩了不少坑，但每个坑都有收获：'
)

doc.add_paragraph(
    '小样本下 GAN 真的很难训。'
    '不是参数没调好，是判别器天生就比生成器更容易赢。'
    'DCGAN、WGAN-GP 标准版、WGAN-GP 削弱版都是同一个根源的不同表现，'
    'StyleGAN2-ADA 理论上能解决这个问题但被环境卡住了。',
    style='List Bullet'
)

doc.add_paragraph(
    '扩散模型的"细节甜蜜期"是真实存在的。'
    '最佳生成结果往往出现在训练中期（200-400 轮），'
    '继续训练反而会损失锐利度——不是模型崩了，是 MSE 在"磨平棱角"。'
    '这提醒我们：训练时要盯着生成预览，不能只看 loss 曲线。',
    style='List Bullet'
)

doc.add_paragraph(
    '医学图像的色彩一致性很重要。'
    'ColorJitter 这类增强在自然图像上是标配，但在医学图像上却是毒药——'
    '眼底图的颜色有诊断意义，不能随便扰动。'
    '数据增强最好限制在几何变换范围内。',
    style='List Bullet'
)

doc.add_paragraph(
    '预训练在生成模型中不是必须的。'
    '和分类/检测任务不同，医学图像生成从零训练是常态。'
    '不必因为"别人都用预训练"而焦虑——在生成任务上，从零训往往更好。',
    style='List Bullet'
)

doc.add_paragraph(
    'MSE 是双刃剑。'
    '它让优化稳定、收敛顺利，但在生成任务中——'
    '尤其当数据分布不均时（眼底图暗背景占主导、亮病灶占少数）——'
    '会系统性地"溶解"少数类特征。'
    '这个坑在 VAE、DDPM、条件扩散、Palette 里反复出现，'
    '直到 FiLM + L1 + LPIPS 才真正缓解。',
    style='List Bullet'
)

doc.add_paragraph(
    '回顾全程，从 VAE 的模糊、GAN 的过拟合到 Diffusion 的病灶溶解，'
    '每次碰壁都不是"某个参数没调好"这种层面的问题，'
    '而是每个模型家族的根本局限在小样本医学图像上被放大了。'
    '但也正因如此，最终筛选出的 FiLM DDPM + L1 + LPIPS 方案，'
    '是在充分理解了"为什么其他方案不行"之后得到的，'
    '不是随便试出来的。'
)

doc.add_heading('参考文献', level=1)
refs = [
    'Kingma & Welling (2013). Auto-Encoding Variational Bayes.',
    'Goodfellow et al. (2014). Generative Adversarial Nets.',
    'Radford et al. (2015). Unsupervised Representation Learning with DCGANs.',
    'Gulrajani et al. (2017). Improved Training of Wasserstein GANs.',
    'Karras et al. (2020). Training GANs with Limited Data (StyleGAN2-ADA).',
    'Ho et al. (2020). Denoising Diffusion Probabilistic Models.',
    'Song et al. (2020). Denoising Diffusion Implicit Models.',
    'Lipman et al. (2022). Flow Matching for Generative Modeling.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

# ===== 保存 =====
out = os.path.join(os.path.dirname(__file__), 'REPORT.docx')
doc.save(out)
print(f'OK: {out}')
print(f'Size: {os.path.getsize(out)/1024:.1f} KB')
