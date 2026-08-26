# -*- coding: utf-8 -*-
"""生成 REPORT_ML.docx —— 传统机器学习方法生成眼底彩照（Phase C2）。

数据源：eval_data/ 下各模型 JSON（_metrics.json + _fundus_metrics.json + _scores.json）+ 结果图。
参考：REPORT（原版）_with_metrics.docx 的工程叙述风格 + _insert_metrics.py 的表格样式。
运行：python _build_report_ml.py（在 research-report/ 下），输出 REPORT_ML.docx。
不修改任何已有文件。
"""
import json
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
EVAL = os.path.join(ROOT, "eval_data")
ML_FIG = os.path.join(ROOT, "generate_project", "machine_learning", "report", "figures",
                      "trad-ml-vs-deep_comparison.png")
ML_RES = os.path.join(ROOT, "generate_project", "machine_learning", "report", "figures")
DST = os.path.join(HERE, "REPORT_ML.docx")

# ---------------- 数据加载 ----------------
def load(name, suffix):
    p = os.path.join(EVAL, f"{name}{suffix}.json")
    return json.load(open(p, encoding="utf-8")) if os.path.exists(p) else {}

GM = {m: load(m, "_metrics") for m in ["pca", "gmm", "patch", "poisson", "retinex"]}
FM = {m: load(m, "_fundus_metrics") for m in ["pca", "gmm", "patch", "poisson", "retinex"]}
SCORES = load("_scores", "")

def score_of(model):
    return SCORES.get("scores", {}).get(model)

def dims_of(model):
    return SCORES.get("dims_scores", {}).get(model, {})

# ---------------- docx 工具 ----------------
doc = Document()

def set_zh(style, zh, en=None, size=None, bold=None, color=None):
    style.font.name = en or zh
    style._element.rPr.rFonts.set(qn("w:eastAsia"), zh)
    if size: style.font.size = Pt(size)
    if bold is not None: style.font.bold = bold
    if color: style.font.color.rgb = RGBColor(*color)

set_zh(doc.styles["Normal"], "等线", "Calibri", 11)
for s in ["Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
    set_zh(doc.styles[s], "微软雅黑", "Microsoft YaHei")

def H1(text): doc.add_heading(text, level=1)
def H2(text): doc.add_heading(text, level=2)
def H3(text): doc.add_heading(text, level=3)

def P(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix + "："); r.bold = True
    p.add_run(text)
    return p

def CODE(text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
    return

def IMG(path, width_in=6.0, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs: r.font.size = Pt(9)
    else:
        P(f"[图缺失: {path}]")

def make_table(header, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(header), style="Table Grid")
    if widths:
        t.autofit = False
        total = sum(widths)
        tblPr = t._tbl.tblPr
        tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), str(total)); tw.set(qn("w:type"), "dxa")
        tblPr.append(tw)
        lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
        grid = t._tbl.find(qn("w:tblGrid"))
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
            gc.set(qn("w:w"), str(w))
        for j, w in enumerate(widths):
            for cell in t.columns[j].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcw = OxmlElement("w:tcW"); tcw.set(qn("w:w"), str(w)); tcw.set(qn("w:type"), "dxa")
                tcPr.append(tcw)
    for j, h in enumerate(header):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            run = cells[j].paragraphs[0].add_run(str(v))
            run.font.size = Pt(font_size)
    return t

def BULLET(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text); return p

def sixdim_table(model_label, model_key):
    """单个方法 + 深度最佳的六维评分对比表。"""
    labels = ["D1 病灶保留", "D2 抗识破+分布", "D3 多样性/质量", "D4 血管结构", "D5 颜色", "D6 记忆风险"]
    d = dims_of(model_key)
    d0 = dims_of("film_l1lpips")
    rows = [[l, f"{d0.get(l, float('nan')):.2f}", f"{d.get(l, float('nan')):.2f}"] for l in labels]
    rows.append(["总分(0-100)", f"{score_of('film_l1lpips'):.1f}", f"{score_of(model_key):.1f}"])
    make_table(["评分维度", "FiLM+LPIPS(深度最佳)", model_label], rows,
               widths=[4200, 3000, 3000])

# ==================================================================
# 封面
# ==================================================================
t = doc.add_paragraph("REPORT_ML：传统机器学习方法生成眼底彩照（Phase C2）")
t.style = doc.styles["Title"]
sub = doc.add_paragraph("小样本重度 DR 图生成的边界与组合潜力 —— 面向小白")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
meta = doc.add_paragraph("2026-08-26 ｜ 数据：330 张 KW-IV 重度眼底彩照 ｜ 方法：PCA / GMM / 补丁拼接（基线）+ 泊松病变重排 / Retinex 光照交换（组合）")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)

# ==================================================================
# 1. 数据集
# ==================================================================
H1("1. 数据集")
P("本报告用 330 张严重糖尿病视网膜病变（DR，Keith-Wagener IV 级）眼底彩照，试了三类传统机器学习方法：PCA 线性重建、GMM 混合采样、补丁拼接。它们都没有神经网络，不反向传播、不梯度训练，是纯传统 ML，回应老师 Bug 2 里\"方法不限于深度学习\"的要求。")
P("这批图跟 Phase A/B 深度方法用同一份数据、同一口径：统一缩放到 128×128 RGB，评估时拿 330 张真实图做对照。")
P("有一件事对传统 ML 既是机会也是坑：KW IV 级意味着每张图都病灶密集，出血、渗出、微动脉瘤到处都是。好处是它自带一座\"病变模板库\"，后面泊松病变重排就靠它；坏处是\"重度\"让数据分布高度非高斯，正好打在参数化模型最弱的地方。")
make_table(["项目", "数值", "说明"], [
    ["样本数", "330", "全部 KW-IV 重度 DR，与深度方法同源"],
    ["分辨率", "128×128 RGB", "统一缩放口径（与 Phase A/B 一致）"],
    ["评估对照", "330 张真实图", "eval_data/real/"],
    ["生成量", "基线 300 张 / 组合 60 张", "旧基线 300 张；组合按 2026-08-18 样本量新规取 60 张"],
], widths=[2500, 2200, 5300])

# ==================================================================
# 2. 评估指标与评分标准
# ==================================================================
H1("2. 评估指标与评分标准")
H2("2.1 为什么用同一套评估体系")
P(f"Phase A/B 已经为 6 个深度方法建好了一套评估体系（通用层 + 专用层 + 六维门控综合评分）。传统 ML 方法复用同一套，才能和深度方法在同一个标尺上横向对比，\"补丁拼接 {score_of('patch'):.1f} 分 vs FiLM 扩散 {score_of('film_l1lpips'):.1f} 分\"这类结论才有意义。否则各说各话，报告没有说服力。")

H2("2.2 通用层指标（对标文献）")
P("10 项通用指标，方向：FID/KID/MMD/MS-SSIM/LPIPS 越低越好；IS/Precision/Recall/Density/Coverage 越高越好；1-NN 越接近 0.5 越好。FID/KID 用 ImageNet 预训练 Inception 特征，128×128 输入。")
make_table(["指标", "含义（小白版）", "方向"], [
    ["FID", "生成图分布离真实图有多远", "↓ 低"],
    ["KID", "FID 的无偏版本，小样本更稳", "↓ 低"],
    ["1-NN", "真假混合最近邻能否分开（0.5=分不开）", "→ 0.5"],
    ["MS-SSIM", "内部多样性（两两相似度）", "↓ 低"],
    ["LPIPS-NN", "最近邻感知距离（越大越不复制）", "↓ 低"],
    ["IS", "类内清晰+类间多样", "↑ 高"],
], widths=[1500, 6200, 1300])

H2("2.3 专用层自设计指标（回应 Bug 1 的\"自设计\"）")
P("针对眼底图设计：病灶保留（出血/渗出）用 Wasserstein 距离 + 保留率（≈1 理想，>1 偏多 <1 偏少）；血管结构用占比分布距离；抗识破用 C2ST（训练小 CNN 区分真假，AUC 越低越难被识破）；记忆风险用最近邻 SSIM（>0.85 占比 = 接近复制比例）；图像质量用 BRISQUE（真实眼底 = 3.69，仅参考）。")

H2("2.4 六维门控综合评分（0-100）")
P("六个维度加权：D1 病灶保留 0.30、D2 抗识破+分布 0.25、D3 多样性/质量 0.20、D4 血管结构 0.10、D5 颜色 0.08、D6 记忆风险 0.07。门控：D2 是\"现实主义\"闸门，D2 太低则总分被扣，防\"糊但指标好看\"的假象。")

H2("2.5 复制检测三件套（质检门控）")
P("330 张小样本处于记忆/复制高风险区（文献显示 GLO 在 128 张就开始记忆训练集）。每个方法生成后必须过三件套：最近邻 NN-SSIM/LPIPS 距离分布、全图复制率、补丁重复率。任一超标，停下重做。")

H2("2.6 人眼评分说明")
P("综合评分是客观指标。人眼评分（人对\"像不像眼底图\"的主观打分）对 ML 方法标注为【待人工评估】，本报告不虚造人工分。各方法章给了效果图和客观分，人眼分等人工复核后补录。")

H2("2.7 踩坑教训：C2ST 必须跑")
P("报告早期按\"传统 ML 无条件图跳过 C2ST\"执行，补丁拼接打出 62.9 分的伪高分，逼近当年深度最佳 70.9（min-max 相对分，模型库扩大后已整体平移，现版各分数见第 8 章）。根因是评分口径不均：深度模型的 D2 维度含 C2ST（权重 0.40），ML 基线没有，两边 D2 在不同指标集上对比。补跑 C2ST 后补丁大幅回落，GMM/PCA 同步重估，才全部同口径。这条教训写进了执行计划：C2ST 必跑，不能跳过。")

# ==================================================================
# 3. 传统 ML 方法谱系与可行性
# ==================================================================
H1("3. 传统 ML 方法谱系与可行性")
H2("3.1 一句话：18 法分三档")
P("文献调研盘点了 18 种传统 ML 图像生成方法，按\"在 330 张重度眼底图上可行的程度\"分三档：")
make_table(["档位", "代表方法", "在眼底图上的判断"], [
    ["高", "Image Quilting 拼布、眼底补丁合成（Fiorini 2014 / Magnusson 2021）、稀疏字典、Poisson 融合、颜色/直方图匹配", "非参数\"重排真实块\"路线，保纹理真实，是唯一能保\"相似\"的路线"],
    ["中", "COCA（Copula PCA 修正）、GMM 紧凑表示、AAM/SSM 结构插值、示例式修补", "有改进空间，但离直接生成仍有距离"],
    ["低", "纯 PCA/PPCA/FA、KDE、FRAME、纯 MRF", "参数化/密度模型，必败于非高斯+强结构（病灶被平均掉）"],
], widths=[1100, 5200, 3700])

H2("3.2 核心规律：分布假设决定能力边界")
P("读完 18 法，最深的一条规律是：传统 ML 生成器的能力边界，由它假设的数据分布形状决定。")
BULLET("参数化模型（PCA/GMM/密度）假设数据是高斯/线性的，而眼底图是强非高斯（黑背景/亮视盘/暗病灶）+ 强结构（血管走向）。线性假设把一切拉向\"平均脸\"，病灶被平均掉，这是它们必败的根源。")
BULLET("非参数补丁法不假设全局分布，直接重排真实纹理块，天然保\"相似\"；但它不产新结构，缺全局布局。")
P("结论：单一传统 ML 方法不行，但\"重排真实块 + 另一层传统方法补全局结构\"的组合管线可行。这是后面章节的出发点。")

H2("3.3 为什么选这三个做基线")
P("三个基线恰好覆盖方法谱系的两端和一个特例：PCA 是参数化线性代表，GMM 是参数化密度代表，补丁拼接是非参数/无模型特例，连参数都没有。把它们跑出来，同时验证\"参数化必败\"和\"非参数保相似缺结构\"两条规律。")

H2("3.4 生成 vs 增广：C2 的价值定位")
P("文献证据（03 篇）显示，经典增广（旋转/mixup）对分类器判别力的提升普遍不输甚至超过 GAN/扩散生成图扩充。所以本报告不争\"传统 ML 生成比增广强\"，而说：传统 ML 组合能产出经典增广做不到的新病灶组合。旋转/mixup 不产生新病灶布局，泊松病变重排能。C2 的价值锚定在\"新组合\"，不在\"更强的判别力\"。")

H2("3.5 七方法效果总览")
P("真实原图、深度最佳、两个组合方法、三个基线的并排对比，每行 3 个样本。一眼可辨：真实图清晰，扩散最接近真实，PCA/GMM 过平滑成\"模糊团\"，补丁是清晰的四块拼贴，泊松/Retinex 保留了真实纹理但带上了新变化。")
IMG(ML_FIG, 6.3, "图 1：真实 vs 深度最佳(FiLM 扩散) vs 组合(泊松病变重排/Retinex 光照交换) vs 三个传统 ML 基线生成效果对比（现版综合分）")

# ==================================================================
# 4. PCA 线性重建（基线 1）
# ==================================================================
H1("4. PCA 线性重建（基线 1）")

H2("4.1 是什么")
P("PCA（主成分分析）是最经典的传统 ML 线性方法。把每张 128×128 图拉成 49152 个数字的向量，330 张图就是 330 个点，挤在 49152 维空间里。PCA 给这一堆点找几个\"最重要的方向\"（主成分），最后只留 64 个。生成一张图 = 给这 64 个方向各配一个权值，随机摇一组，反投回像素。")
P("这里的\"重要\"有明确的数学定义：主成分按方差排序。第 1 个方向是这堆点最散开的方向，第 2 个是剩余差异里最散开、且跟第 1 个垂直的方向，依此类推。64 个主成分累计保留 96% 的方差。方差就是图与图之间的差异，留 96% 意味着砍掉 4%，而这 4% 恰恰是锐利的病灶边缘、细小的血管分叉这类高频细节。")
P("有个小白最容易误解的点：PCA 没有训练轮数。它是一次矩阵特征分解的解析解，一步到位，不存在多训几轮变好的可能。")

H2("4.2 怎么做（略写）")
CODE("python scripts/pca_gen.py --num_images 300 --seed 42\n# PCA 拟合 64 维（保留 96% 方差）→ 潜空间随机采样 → 反投影 → clip 0-255")

H2("4.3 为什么结果必然过平滑")
P("压缩到 64 维，模型只有 64 个自由度去表达整张图。主成分本身是\"全体图的平均方向\"，天生平滑。生成图就是这些平滑方向的加权和，细节早在降维那一步就被当噪声扔了，无论怎么算都找不回来。")
P("打个粗糙的比方：给一台钢琴装 64 个旋钮去复刻一个乐队。能拧出大致的旋律走向（平均脸），但每个乐手的指法细节（病灶、血管）是旋钮表达不出来的。")

H2("4.4 结果")
P(f"总分 {score_of('pca'):.1f} 分（完成时 20.0；min-max 相对分随模型库扩大平移，下同），FID 297.1 全组最差（深度最佳 178.8）。1-NN=1.000 完美可分，IS=1.58 多样性几乎为零。病灶保留率：出血 1.57 偏多、渗出 0.56 偏少，连病灶比例都没对齐。记忆检测通过（接近复制占比 0）。人眼评分【待人工评估】。")
sixdim_table("PCA(基线1)", "pca")
IMG(os.path.join(ML_RES, "result_pca.png"), 4.0, "图 3：PCA 生成结果 6 张（注意整体模糊、病灶消失）")

H2("4.5 致命缺点")
BULLET("过平滑：线性假设把病灶平均掉，没有任何锐利细节")
BULLET("无新病灶形态：所有输出只是\"平均脸\"的随机抖动")
BULLET("无全局结构：血管走向、视盘位置全靠碰运气")

H2("4.6 下一步")
P("PCA 的失败证实参数化线性方法不可行，但它的潜空间还能用。COCA（高斯 copula 修正）能改善非高斯伪影；更重要的是 PCA 可以作为结构层（血管/布局）的紧凑表示，跟纹理层组合。单个不行，组合是出路。")

# ==================================================================
# 5. GMM 混合采样（基线 2）
# ==================================================================
H1("5. GMM 混合采样（基线 2）")

H2("5.1 是什么")
P("GMM（高斯混合模型）是概率密度估计：先用 PCA 把 330 张图压到 64 维，再用 16 个\"高斯团\"盖住这批点的分布。每个团是一个多维高斯（钟形），有自己的中心、胖瘦、朝向。生成 = 先随机挑一个团，再从团里随机采一个点，反投回像素。")
P("参数靠 EM 算法迭代学出来。EM 交替做两件事：按当前的团参数估计每个点更可能属于哪个团（E 步），再按归属重新估计团的参数（M 步），反复到稳定。实测给 10/100/1000 次迭代上限，它 2 次就收敛，BIC 一点没变。所以过平滑不是\"没训够\"，是生成机制本身的问题。", "怎么学的")

H2("5.2 怎么做（略写）")
CODE("python scripts/gmm_gen.py --num_images 300 --seed 42\n# PCA 降维 64 → GMM(16 分量) 拟合 → 混合分布采样 → 反投影")

H2("5.3 为什么结果必然过平滑")
P("高斯分布的形状是光滑的钟形。从一个高斯团里采样，点永远落在中心附近，不可能跑到团外的尖角区域。病灶边缘是\"尖\"的，高斯是\"圆\"的，圆表达不了尖，血管分叉同理。模型想表达也表达不出来。")
P("GMM 比 PCA 强在能刻画多模态：16 个团可以各自对应一种典型长相，比如两种血管走向各占一个峰，所以能出大致形状和轮廓。但细节依然被抹平。")
P("类比：一群人的身高可以用几个高斯峰描述，一个峰代表成年人、一个代表儿童。但身高分布再准也描述不了每个人的头发。眼底图的问题是，病灶细节恰好是\"头发\"这类被分布抹平的东西。")

H2("5.4 结果")
P(f"总分 {score_of('gmm'):.1f} 分（完成时 40.2），FID 233.8，1-NN=1.000。比 PCA 好：能出大致形状和轮廓，病灶保留率也更接近 1（出血 1.43、渗出 0.71）。但病灶仍模糊，放大即露馅。记忆检测通过。人眼评分【待人工评估】。")
sixdim_table("GMM(基线2)", "gmm")
IMG(os.path.join(ML_RES, "result_gmm.png"), 4.0, "图 4：GMM 生成结果 6 张（轮廓接近真实但细节仍糊）")

H2("5.5 致命缺点")
BULLET("过平滑：高斯分量的光滑性决定了病灶边缘模糊")
BULLET("无新病灶形态：采样只会在已有分布附近抖动，造不出新的病变组合")
BULLET("对 330 张重度图，多模态优势发挥有限，分布本身太集中")

H2("5.6 下一步")
P("GMM 比 PCA 稍好，但天花板相同。它的价值同样是\"层次之一\"：GMM 可以在结构层对血管/病灶布局建模，跟纹理层组合。组合管线再次是答案。")

# ==================================================================
# 6. 补丁拼接（基线 3）
# ==================================================================
H1("6. 补丁拼接（基线 3）")

H2("6.1 是什么")
P("补丁拼接是无模型方法，连参数都没有。从 330 张真实图里随机裁 64×64 的真实小块，4 块拼成一张 2×2 马赛克。它不压缩、不平均、不生成，只是把真实像素换个位置。这是 Image Quilting（Efros & Freeman 2001）的最朴素版本，文献里拼布法在眼底分割上验证过（Fiorini 2014 / Magnusson 2021）。")

H2("6.2 怎么做（略写）")
CODE("python scripts/patch_gen.py --num_images 300 --seed 42\n# 随机选图 + 随机起点裁 64×64 真实块 → 4 块填入 2×2 → 保存")

H2("6.3 为什么它\"保相似但不生成\"")
P("每一步都是真实像素的直接搬移，纹理、颜色、病灶比例全部继承真实图，所以它\"像\"。但它不做结构：4 块之间没有连续性，接缝处血管、视盘、病灶布局对不上。输出是清晰但机械的拼贴。")
P("类比：把四张真照片剪开拼成一张。每一块都真实，但整张图不是任何一只真实的眼睛。")

H2("6.4 结果（含门控发现）")
P(f"总分 {score_of('patch'):.1f} 分（完成时 42.1），FID 159.2 三个 ML 里最低，甚至比部分深度模型还低，1-NN=0.995。病灶保留率最接近 1（出血 0.94、渗出 1.16），血管 Wasserstein 最小，因为像素本来就是真的。C2ST=1.000：真实像素拼接被小 CNN 一眼识破。记忆检测通过。人眼评分【待人工评估】。")
sixdim_table("补丁拼接(基线3)", "patch")
P("本方法最早得 62.9 分伪高分，排查后确认是 C2ST 缺失导致的口径不均（见 2.7），补跑后回落到正常水平。这个\"伪高分事件\"本身有价值：它暴露了评分体系的盲区，凡是用真实像素的方法，病灶/血管/颜色三维必然高分。这里如实记录，作为评估体系改进的依据。")
IMG(os.path.join(ML_RES, "result_patch.png"), 4.0, "图 5：补丁拼接结果 6 张（清晰但 4 块拼贴、接缝可见）")

H2("6.5 致命缺点")
BULLET("无意义机械重排：每张都不同，但没有医学意义的新组合")
BULLET("接缝伪影：块间不做平滑/对齐")
BULLET("局部复制风险：每一像素都来自真实图，需复制检测盯防")

H2("6.6 下一步")
P("补丁法证明了非参数能保住纹理真实。它缺的全局结构，恰好可以由另一层传统方法补上：血管骨架生成 + 病灶布局规划。这就是组合管线的思路：泊松病变重排（保留真实病灶、加上宿主图结构）和 Retinex 光照交换（保留结构、交换光照），都比纯拼接多一层语义。")

# ==================================================================
# 7. 高价值组合：泊松病变重排 + Retinex 光照交换（已实现）
# ==================================================================
H1("7. 高价值组合：泊松病变重排 + Retinex 光照交换（已实现）")
P("三个基线的结论指向同一个方向：单一传统 ML 不行，但组合可行。文献里最有支撑的两条组合（均有眼底专版验证）已按六段式跑完：泊松病变重排（EX-004）与 Retinex 光照交换（EX-005）。它们都产出了三个基线做不到的东西：保留真实纹理的同时引入新变化。")

H2("7.1 泊松病变重排（EX-004，80.8 分）")
H3("7.1.1 是什么")
P("330 张全是重度 DR，等于自带一座病变模板库。泊松重排从图 A 抠一簇渗出、图 B 抠一块出血，用泊松编辑（cv2.seamlessClone 的 NORMAL_CLONE）无缝贴到宿主图 D 上，再叠一层全局旋转和色彩微调，产出\"新病灶组合\"的严重 DR 图。")
P("这里有个省事的设计：方法不需要知道哪块是病灶。随机椭圆贴片簇（半径 60-150px @512）就行，重度图到处是病灶，随便搬一块就是病灶样内容。", "为什么不用病灶检测")

H3("7.1.2 怎么做（略写）")
CODE("python scripts/poisson_gen.py --num_images 300 --seed 42 --max_patches 10 --patch_radius 100 --rot 15 --color_jit 0.18")

H3("7.1.3 为什么有用")
P("泊松编辑不在像素值上直接做，而是操作梯度，也就是相邻像素的差值。它把供体病灶区域的梯度整体搬过来，再要求拼接边界处梯度跟宿主图连续，解一个线性方程（泊松方程）。梯度连续意味着颜色过渡自然，接缝肉眼看不见。直接 copy-paste 像素会在边界留一圈色差，泊松编辑专门消掉这个。", "机制")
P("它生成的是经典增广做不到的东西。旋转、裁剪、mixup 都只在\"已有的那张图\"上做手脚，泊松是把两张图的病灶重新组合。文献证据（Yu 2021 BOE）显示，这种物体级重排对 DR 筛查的提升超过过采样/裁剪/旋转。", "为什么是\"新组合\"")

H3("7.1.4 结果")
P("总分 80.8（完成时 87.5；min-max 相对分随模型库扩大平移，下同），全场第 1。六维 D1 0.97 / D2 0.87 / D3 0.91 / D4 1.00 / D5 1.00 / D6 0.00。复制率 0.3%。FID 58.1（深度最佳 178.8），C2ST 0.804 全场第 2 低，BRISQUE 3.60≈真实 3.69，病灶保留出血 1.10/渗出 0.97，血管 0.070 vs 真实 0.077。人眼评分：用户确认\"几乎看不出来是生成的\"。")
sixdim_table("泊松病变重排(EX-004)", "poisson")
IMG(os.path.join(ML_RES, "result_poisson.png"), 3.2, "图 6：泊松病变重排结果 4 组（donor｜底图｜生成 对比，可见病灶被搬运重组）")

H3("7.1.5 致命缺点")
BULLET("不生成新解剖结构：血管/视盘/病灶布局仍是宿主图的，只是换了位置/数量")
BULLET("高分部分源于底图复用：病灶/血管/颜色三维是真实像素的平凡胜利，如实标注")
BULLET("D6 记忆 0.00：最接近真实集，但也意味着离底图最近，需复制检测盯防（复制率 0.3% 已可控）")

H3("7.1.6 下一步")
P("泊松验证了\"新病灶组合\"假设成立：重排比纯拼接多一层语义，保留宿主图结构的同时搬来真实病灶。它进入 Phase D 下游分类器验证候选，重点检验生成的重度图能否提升分类器对重度类的 Recall。")

H2("7.2 Retinex 光照交换（EX-005，72.5 分）")
H3("7.2.1 是什么")
P("把人眼看到的图拆成两层：光照 L × 反射 R。光照是打在视网膜上的光，决定明暗和色温；反射是物体本身的图案，决定血管、视盘、病灶长什么样。Retinex 把图取对数，在 log 域用一个很大的高斯核做低通，平滑出来的就是光照 L，原图减掉它就是反射 R。")
P("在 log 域做是因为光照乘反射的关系取对数后变成加法，分离起来干净。大核（41px @512）保证光照只剩平滑渐变，所有锐利结构都留在反射里。", "为什么 log 域 + 大核")
P("交换光照按公式重组：新图 = R_A · L_A^(1-α) · L_B^α。α=0 完美还原底图 A，α=1 换成供体 B 的纯光照，0<α<1 是 log 域线性插值。每张随机取 α∈[0.6,0.9]，保证光照变化明显。", "交换公式")
P("早期版本公式写错过：new = R_A · L_B^α 只乘供体光照，把幅度也压缩了（90^0.8≈36 而非 90），输出近全黑。改成 log 域插值后，α=0 还原底图（误差 0），α=0.8 亮度 144 vs 底图 150。", "踩过的坑")

H3("7.2.2 怎么做（略写）")
CODE("python scripts/retinex_gen.py --num_images 60 --seed 42 --kernel 41 --alpha_lo 0.6 --alpha_hi 0.9 --mem_ref ../../eval_data/real --mem_thr 0.85 --max_retry 20\n# Retinex 分解 → 跨图光照 α 插值 → 重组 → 近复制过滤（NN-SSIM>0.85 重采样）")

H3("7.2.3 为什么有用")
P("因为光照和反射解耦了：换光照只改\"看起来的调子\"，诊断语义（血管、病灶、视盘）几乎不动。这是深度无条件生成器很难显式控制的维度，你没法告诉扩散模型\"只把光照换成另一张图的\"。Retinex 用显式规则就做到了。", "机制")
P("文献（Zhang 2022 CBM）用无监督 Retinex 增强光照层、保反射层，眼底分割提升 9.6% Dice。本脚本是它的经典版改编：不做网络，直接跨图交换光照。", "文献")

H3("7.2.4 结果")
P("总分 72.5，全场第 2。六维 D1 0.94 / D2 0.82 / D3 0.72 / D4 0.92 / D5 0.92 / D6 0.10。复制率 0%（方案 3 过滤后归零）。C2ST AUC 0.744 全场最佳，最难被识破；FID 117.9 / KID 0.056 全场最低（底图复用所致）；BRISQUE 3.70≈真实 3.69。病灶保留出血 1.06/渗出 0.44（渗出偏低）。亮度中位 77（真实 61，部分图偏暗）。人眼评分【待人工评估】。")
sixdim_table("Retinex光照交换(EX-005)", "retinex")
IMG(os.path.join(ML_RES, "result_retinex.png"), 3.2, "图 7：Retinex 光照交换结果 4 组（donor｜底图｜生成 对比，结构继承底图、光照来自供体）")

H3("7.2.5 致命缺点")
BULLET("不生成新病灶形态：结构完全继承底图，病灶组合/布局不变")
BULLET("平均 NN-SSIM 0.439 偏高 → D6 记忆 0.10 弱项（光照交换天然保留底图结构）")
BULLET("亮度受供体光照影响：部分图偏暗/偏亮")

H3("7.2.6 下一步")
P("Retinex 验证了\"安全多样化\"假设：抗识破全场最佳、复制归零，适合做\"相似但不相同\"的多样性放大器。同样进入 Phase D 下游验证候选。")

H2("7.3 可选：血管骨架 + 纹理拼布管线（后续方向）")
P("用传统方法（ASM/Bonaldi 或过程化 CCO/DLA）生成全新血管骨架，沿骨架拼真实纹理，inpainting 补缝。结构全新 + 纹理真实，最接近\"真生成\"，但工程量最高。（Fiorini 2014 / Magnusson 2021 已在眼底验证并提升分割 SOTA）")

H2("7.4 为什么这两个组合值得做")
P("深度无条件生成器难以显式控制\"病灶种类/密度/位置\"和光照的组合；泊松/Retinex 用显式规则做到，泊松造新病灶组合、Retinex 换光照风格。这是本报告的核心论点：传统 ML 组合做深度方法做不到的事，作为数据扩充有补充价值。")

# ==================================================================
# 8. 结论
# ==================================================================
H1("8. 结论")
H2("8.1 项目回顾")
P(f"本报告用 330 张严重 DR 眼底彩照，试了三类传统 ML 基线（PCA {score_of('pca'):.1f} / GMM {score_of('gmm'):.1f} / 补丁 {score_of('patch'):.1f}）和两条组合管线（泊松病变重排 {score_of('poisson'):.1f} / Retinex 光照交换 {score_of('retinex'):.1f}），跟 Phase A/B 的 6 个深度方法在同一套评估体系下横向对比（深度最佳 FiLM+LPIPS 当前 {score_of('film_l1lpips'):.1f}）。综合评分是 min-max 相对分数，会随模型库扩大平移，本报告各表用同一版 _scores.json。")

IMG(os.path.join(ROOT, "research-report", "figures", "score_overview.png"), 6.0, "图 2：11 个模型综合评分总览（六维门控 0-100，min-max 相对分）")

H2("8.2 核心发现")
BULLET(f"纯传统 ML 直接生成不可行：PCA/GMM 的过平滑证实参数化方法必败于非高斯+强结构；补丁虽得 {score_of('patch'):.1f}，但那是真实像素的平凡胜利，不是生成。")
BULLET("评分体系盲区被暴露：补丁伪高分（62.9 → 补跑 C2ST 后大幅回落）证明用真实像素的方法必然在病灶/血管/颜色维度高分，跨\"真实重排 vs 生成\"对比时 C2ST 必须跑。")
BULLET("非参数保\"相似\"、参数化保不了：补丁法保纹理真实但缺结构，恰好引出组合管线。")
BULLET("组合管线让\"相似但不相同\"落地：泊松病变重排（80.8）与 Retinex 光照交换（72.5）都大幅超过三个基线，复制率可控（0.3% / 0%）。泊松造新病灶组合、Retinex 换光照风格，做的是深度无条件生成器难以显式控制的事。")

H2("8.3 C2 的价值定位")
P("本报告的价值不在于\"传统 ML 比深度强\"（文献不支持，也不争）。它有两点：一是完整的方法论对照，说明为什么深度生成/组合管线是必要的，回应老师 Bug 2；二是传统 ML 组合能产出深度无条件生成器难以显式控制的新病灶组合，作为数据扩充有补充价值。")

H2("8.4 局限")
BULLET("无真实病灶 mask：病灶保留率基于像素统计估计，非专家标注")
BULLET("330 张小样本：记忆/复制风险高，评估必须患者级划分")
BULLET("128×128 分辨率偏低；BRISQUE 在自然图像上训练，眼底仅参考")

H2("8.5 下一步")
P("阶段 2 已完成：泊松病变重排（80.8）验证\"新病灶组合\"、Retinex 光照交换（72.5）验证\"安全多样化\"假设均成立。下一步进入 Phase D 下游分类器验证（TSTR/TRTR 协议）：用生成的严重 DR 图扩充 DR 分级分类器训练集，验收\"合成重度图显著提升对应类 Recall/F1/平衡准确率\"。每步实验结果按六段式更新到本报告。")

# ==================================================================
# 参考文献
# ==================================================================
H1("参考文献")
refs = [
    "Efros A A, Freeman W T. Image Quilting for Texture Synthesis and Transfer. SIGGRAPH 2001.",
    "Pérez P, Gangnet M, Blake A. Poisson Image Editing. SIGGRAPH 2003.",
    "Fiorini S, et al. STAG: fundus image synthesis by image quilting (视网膜合成). 2014.",
    "Bonaldi L, et al. Automatic generation of synthetic retinal fundus images (ASM 血管树). 2016.",
    "Magnusson J, et al. DeLTA: fundus segmentation with synthetic data (DRIVE/STARE SOTA). 2021.",
    "Yu Z, et al. BOE: Biomedical image synthesis via object-level rearrangement (泊松病变重排). 2021.",
    "Zhang J, et al. CBM: Color-consistent fundus image synthesis with Retinex (光照交换). 2022.",
    "Egger B, et al. COCA: Copula Eigenfaces for face synthesis (Copula PCA). 2016.",
    "Feng Q, et al. Training data memorization in generative models (ICCV 2021).",
    "Ahamed M, et al. Mixup for diabetic retinopathy classification. 2025.",
    "Beinecke J, et al. Data augmentation in medical imaging: when does it help? 2021.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}] {r}")
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs: run.font.size = Pt(9)

doc.save(DST)
print(f"已生成 {DST}")

# ---- 简单自检 ----
h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
print(f"H1 章数: {len(h1)}")
for h in h1: print("  ", h)
print(f"表格数: {len(doc.tables)}  段落数: {len(doc.paragraphs)}")
